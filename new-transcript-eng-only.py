import os
from docx import Document
import whisper
from dotenv import load_dotenv
import openai
from openai import OpenAI
import re
import time

# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Whisper model
print("[INFO] Loading Whisper model (medium.en)...")
whisper_model = whisper.load_model("medium.en", device="cpu")
print("[INFO] Whisper model loaded successfully!")

# Updated prompt that specifically handles transitional words
FORMATTING_PROMPT = """You are helping to format an English transcript from Nomad Coders Nicolas's lecture into natural subtitle segments.

Your task:
• Take the raw English transcript and break it into natural subtitle chunks (1-2 sentences each)
• Break based on natural speech patterns, emphasis points, and logical pauses
• Don't break in the middle of important concepts or explanations
• Each subtitle segment should be complete and make sense on its own
• Keep the original English text exactly as is - don't modify, correct, or change anything

IMPORTANT RULES for transitional words:
• Words like "okay?", "alright?", "yep", "yes", "perfect", "cool", "right?", "you know?" should ALWAYS stay attached to the previous sentence
• These should NEVER be on their own line
• They provide emphasis and should remain connected to the main thought

• Output each subtitle segment on its own line
• Add a blank line between each subtitle segment for spacing

Example input:
"All right everybody in this video we are going to implement our voice agent okay we're going to turn this agent that is text based into a voice agent alright now I have to tell you that this implementation of voice agents I think is a little bit rough you know"

Example output:
All right everybody, in this video we are going to implement our voice agent, okay?

We're going to turn this agent that is text-based into a voice agent, alright?

Now I have to tell you that this implementation of voice agents I think is a little bit rough, you know?

Keep it natural and flowing like subtitles would appear in a video."""


def post_process_formatted_transcript(formatted_text):
    """
    Post-process the GPT-formatted transcript to fix transitional words
    that ended up on their own lines
    """
    # List of transitional words/phrases that should be combined with previous line
    transitional_patterns = [
        r"^\s*(okay\??)\s*$",
        r"^\s*(alright\??)\s*$",
        r"^\s*(all right\??)\s*$",
        r"^\s*(yep\??)\s*$",
        r"^\s*(yes\??)\s*$",
        r"^\s*(yeah\??)\s*$",
        r"^\s*(perfect\??)\s*$",
        r"^\s*(cool\??)\s*$",
        r"^\s*(right\??)\s*$",
        r"^\s*(you know\??)\s*$",
        r"^\s*(fantastic\??)\s*$",
        r"^\s*(great\??)\s*$",
        r"^\s*(nice\??)\s*$",
        r"^\s*(good\??)\s*$",
        r"^\s*(exactly\??)\s*$",
        r"^\s*(totally\??)\s*$",
        r"^\s*(absolutely\??)\s*$",
        r"^\s*(sure\??)\s*$",
        r"^\s*(got it\??)\s*$",
    ]

    lines = formatted_text.split("\n")
    processed_lines = []
    i = 0

    while i < len(lines):
        current_line = lines[i].strip()

        # Check if current line is a transitional word
        is_transitional = False
        for pattern in transitional_patterns:
            if re.match(pattern, current_line, re.IGNORECASE):
                is_transitional = True
                break

        if is_transitional and processed_lines:
            # Find the last non-empty line to attach this to
            last_content_idx = len(processed_lines) - 1
            while (
                last_content_idx >= 0 and not processed_lines[last_content_idx].strip()
            ):
                last_content_idx -= 1

            if last_content_idx >= 0:
                # Remove any trailing punctuation from the previous line
                prev_line = processed_lines[last_content_idx].rstrip(".,!?")
                # Add the transitional word with appropriate punctuation
                processed_lines[last_content_idx] = (
                    f"{prev_line}, {current_line.lower()}"
                )
                # Skip adding the current line separately
                i += 1
                continue

        processed_lines.append(lines[i])
        i += 1

    return "\n".join(processed_lines)


def format_transcript_with_gpt(transcript_text):
    """
    Format transcript into natural subtitle segments using GPT
    """
    try:
        print("[INFO] Formatting transcript into subtitle segments...")

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": FORMATTING_PROMPT},
                {"role": "user", "content": transcript_text},
            ],
            temperature=0.1,  # Very low for consistent formatting
            max_tokens=15000,
        )

        formatted_transcript = response.choices[0].message.content

        # Post-process to fix transitional words
        formatted_transcript = post_process_formatted_transcript(formatted_transcript)

        print("[INFO] Transcript formatting complete.")
        return formatted_transcript

    except Exception as e:
        print(f"[ERROR] Formatting failed: {e}")
        # Fallback: basic sentence splitting with transitional word handling
        return format_transcript_fallback(transcript_text)


def format_transcript_fallback(transcript_text):
    """
    Fallback formatting method that handles transitional words properly
    """
    # Split into sentences but be smart about it
    sentences = re.split(r"(?<=[.!?])\s+", transcript_text)

    transitional_words = [
        "okay",
        "alright",
        "all right",
        "yep",
        "yes",
        "yeah",
        "perfect",
        "cool",
        "right",
        "you know",
        "fantastic",
        "great",
        "nice",
        "good",
        "exactly",
        "totally",
        "absolutely",
        "sure",
        "got it",
    ]

    formatted_segments = []
    current_segment = ""

    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue

        # Check if this sentence starts with a transitional word
        first_words = sentence.lower().split()[:2]  # Check first 1-2 words
        is_transitional_start = any(
            sentence.lower().startswith(word.lower())
            or " ".join(first_words).startswith(word.lower())
            for word in transitional_words
        )

        if current_segment and (
            len(current_segment) + len(sentence) > 150
            or (not is_transitional_start and len(current_segment) > 50)
        ):
            # Start new segment
            formatted_segments.append(current_segment.strip())
            current_segment = sentence
        else:
            # Add to current segment
            if current_segment:
                current_segment += " " + sentence
            else:
                current_segment = sentence

    # Add the last segment
    if current_segment.strip():
        formatted_segments.append(current_segment.strip())

    # Join with double line breaks for spacing
    return "\n\n".join(formatted_segments)


def transcribe_and_format_audio(file_path, output_directory):
    """
    Main function that transcribes and formats transcript only
    """
    # Transcribe audio
    print(f"[INFO] Transcribing '{file_path}'...")
    result = whisper_model.transcribe(file_path, word_timestamps=True)

    # Get full transcript text
    full_transcript = result["text"].strip()

    if not full_transcript:
        print("[WARNING] No transcription found.")
        return ""

    print(f"[INFO] Transcription length: {len(full_transcript)} characters")

    # Format the transcript into subtitle segments
    formatted_content = format_transcript_with_gpt(full_transcript)

    return formatted_content


def save_transcript_to_docx(content, output_file):
    """
    Save formatted transcript to DOCX
    """
    print(f"[INFO] Saving transcript to '{output_file}'...")
    doc = Document()

    # Split content by lines and handle spacing
    lines = content.split("\n")

    for line in lines:
        if line.strip():
            # Add paragraph with content
            doc.add_paragraph(line.strip())
        else:
            # Add empty paragraph for spacing
            doc.add_paragraph()

    doc.save(output_file)
    print(f"[INFO] Transcript saved successfully!")


def save_transcript_to_txt(content, output_file):
    """
    Save formatted transcript to TXT file as well
    """
    print(f"[INFO] Saving transcript to '{output_file}'...")
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[INFO] TXT file saved successfully!")


# Main execution
if __name__ == "__main__":
    directory = r"C:\Users\Jonathan\Documents\kbdocs\text_result"
    media_files = [f for f in os.listdir(directory) if f.endswith((".mp4", ".mp3"))]

    if not media_files:
        print("[WARNING] No media files found in the directory.")
    else:
        for media_file in media_files:
            audio_file_path = os.path.join(directory, media_file)
            print(f"[INFO] Processing file: {audio_file_path}")

            # Get formatted transcript
            formatted_transcript = transcribe_and_format_audio(
                audio_file_path, directory
            )

            if formatted_transcript:
                print("[INFO] Transcription and formatting completed!")

                # Save to DOCX
                base_name = os.path.splitext(media_file)[0]
                docx_output = os.path.join(directory, f"{base_name}_eng-only.docx")
                save_transcript_to_docx(formatted_transcript, docx_output)

                # Also save to TXT for easy reading
                # txt_output = os.path.join(directory, f"{base_name}_transcript.txt")
                # save_transcript_to_txt(formatted_transcript, txt_output)

                print(f"[INFO] Files saved:")
                print(f"  - DOCX: {docx_output}")
                # print(f"  - TXT: {txt_output}")
            else:
                print("[WARNING] No transcript content to save.")
