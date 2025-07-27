import os
import logging
from pathlib import Path
from typing import List, Optional
from pydub import AudioSegment
import whisper
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
import textwrap
from datetime import datetime
import argparse
import json

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.FileHandler("transcription.log"), logging.StreamHandler()],
)
logger = logging.getLogger(__name__)


class AudioTranscriber:
    def __init__(self, model_name: str = "medium.en", device: str = "cpu"):
        """Initialize the transcriber with specified model and device."""
        self.model_name = model_name
        self.device = device
        self.model = None
        self.supported_formats = {
            ".mp3",
            ".mp4",
            ".wav",
            ".m4a",
            ".flac",
            ".ogg",
            ".webm",
        }

    def load_model(self):
        """Load the Whisper model."""
        if self.model is None:
            logger.info(
                f"Loading Whisper model ({self.model_name}) on {self.device}..."
            )
            try:
                self.model = whisper.load_model(self.model_name, device=self.device)
                logger.info("Whisper model loaded successfully!")
            except Exception as e:
                logger.error(f"Failed to load model: {e}")
                raise

    def format_transcription(self, transcription: str, line_length: int = 90) -> str:
        """
        Format transcription text for better readability.

        Args:
            transcription: Raw transcription text
            line_length: Maximum line length for wrapping

        Returns:
            Formatted transcription text
        """
        keywords_to_combine = [
            "Yes",
            "Yep",
            "Yeah",
            "Perfect",
            "Okay",
            "Sure",
            "Right",
            "Fantastic",
            "All right",
            "Cool",
            "Exactly",
            "Absolutely",
            "No",
            "Nope",
            "Wait",
            "Hold on",
            "Actually",
            "Well",
        ]

        # Clean up the text
        text = transcription.replace("\n", " ").strip()
        if not text:
            return ""

        # Split into sentences more intelligently
        import re

        sentences = re.split(r"(?<=[.!?])\s+", text)

        formatted_text = ""

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            # Check if sentence starts with a keyword
            starts_with_keyword = any(
                sentence.startswith(keyword + " ") or sentence == keyword
                for keyword in keywords_to_combine
            )

            if starts_with_keyword and formatted_text:
                # Combine with previous line
                formatted_text = formatted_text.rstrip() + f" {sentence}\n\n"
            else:
                # Add as new paragraph
                wrapped = textwrap.fill(sentence, width=line_length)
                formatted_text += wrapped + "\n\n"

        return formatted_text.strip()

    def convert_to_wav(self, file_path: Path, output_dir: Path) -> Path:
        """Convert audio file to WAV format."""
        logger.info(f"Converting '{file_path.name}' to WAV format...")

        try:
            audio = AudioSegment.from_file(str(file_path))
            wav_file_path = output_dir / f"{file_path.stem}.wav"

            # Optimize audio settings for Whisper
            audio = audio.set_frame_rate(16000).set_channels(
                1
            )  # Whisper prefers 16kHz mono
            audio.export(str(wav_file_path), format="wav")

            logger.info(f"Conversion complete. Saved as '{wav_file_path.name}'")
            return wav_file_path

        except Exception as e:
            logger.error(f"Error converting {file_path}: {e}")
            raise

    def split_audio(
        self, file_path: Path, chunk_length_minutes: int = 10
    ) -> List[AudioSegment]:
        """Split audio file into chunks."""
        chunk_length_ms = chunk_length_minutes * 60 * 1000
        logger.info(
            f"Splitting '{file_path.name}' into {chunk_length_minutes}-minute chunks..."
        )

        try:
            audio = AudioSegment.from_file(str(file_path))
            chunks = []

            for i in range(0, len(audio), chunk_length_ms):
                chunk = audio[i : i + chunk_length_ms]
                chunks.append(chunk)

            logger.info(f"Created {len(chunks)} chunks")
            return chunks

        except Exception as e:
            logger.error(f"Error splitting audio: {e}")
            raise

    def transcribe_chunk(
        self, chunk: AudioSegment, chunk_index: int, temp_dir: Path
    ) -> str:
        """Transcribe a single audio chunk."""
        temp_file = temp_dir / f"temp_chunk_{chunk_index}.wav"

        try:
            logger.info(f"Transcribing chunk {chunk_index + 1}...")

            # Export chunk to temporary file
            chunk.export(str(temp_file), format="wav")

            # Transcribe with additional options
            result = self.model.transcribe(
                str(temp_file),
                language="en",  # Specify language for better performance
                fp16=False,  # Disable fp16 for CPU compatibility
                verbose=False,  # Reduce verbose output
            )

            # Clean up temp file immediately
            if temp_file.exists():
                temp_file.unlink()

            # Format the transcription
            formatted_text = self.format_transcription(result["text"])
            logger.info(f"Chunk {chunk_index + 1} transcribed successfully")

            return formatted_text

        except Exception as e:
            logger.error(f"Error transcribing chunk {chunk_index + 1}: {e}")
            # Clean up temp file on error
            if temp_file.exists():
                temp_file.unlink()
            raise

    def save_to_docx(self, transcription: str, output_file: Path, audio_file_name: str):
        """Save transcription to a formatted DOCX file."""
        logger.info(f"Saving transcription to '{output_file.name}'...")

        try:
            doc = Document()

            # Add header with metadata
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"Audio Transcription - {audio_file_name}"

            # Add title
            title = doc.add_heading(f"Transcription: {audio_file_name}", 0)

            # Add metadata
            doc.add_heading("Transcription Details", level=1)
            details = doc.add_paragraph()
            details.add_run("File: ").bold = True
            details.add_run(f"{audio_file_name}\n")
            details.add_run("Transcribed: ").bold = True
            details.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            details.add_run("Model: ").bold = True
            details.add_run(f"Whisper {self.model_name}\n")

            # Add transcription content
            doc.add_heading("Transcription", level=1)

            # Split transcription into paragraphs and add them
            paragraphs = transcription.split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())

            doc.save(str(output_file))
            logger.info("Transcription saved successfully!")

        except Exception as e:
            logger.error(f"Error saving DOCX: {e}")
            raise

    def save_metadata(self, metadata: dict, output_file: Path):
        """Save transcription metadata to JSON file."""
        try:
            with open(output_file, "w", encoding="utf-8") as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)
            logger.info(f"Metadata saved to {output_file.name}")
        except Exception as e:
            logger.error(f"Error saving metadata: {e}")

    def process_audio_file(
        self, file_path: Path, output_dir: Path, chunk_length: int = 10
    ) -> dict:
        """Process a single audio file and return metadata."""
        logger.info(f"Processing file: {file_path.name}")
        start_time = datetime.now()

        # Ensure model is loaded
        self.load_model()

        # Create temporary directory for chunks
        temp_dir = output_dir / "temp"
        temp_dir.mkdir(exist_ok=True)

        try:
            # Convert to WAV if needed
            if file_path.suffix.lower() != ".wav":
                wav_file = self.convert_to_wav(file_path, temp_dir)
            else:
                wav_file = file_path

            # Split audio into chunks
            chunks = self.split_audio(wav_file, chunk_length)

            # Transcribe each chunk
            full_transcript = ""
            for i, chunk in enumerate(chunks):
                chunk_transcript = self.transcribe_chunk(chunk, i, temp_dir)
                if chunk_transcript.strip():
                    full_transcript += (
                        f"\n--- Segment {i + 1} ---\n\n{chunk_transcript}\n"
                    )

            # Clean up temporary WAV file if we created it
            if wav_file != file_path and wav_file.exists():
                wav_file.unlink()
                logger.info("Temporary WAV file deleted")

            # Save transcription
            output_docx = output_dir / f"{file_path.stem}_transcript.docx"
            self.save_to_docx(full_transcript, output_docx, file_path.name)

            # Save plain text version
            output_txt = output_dir / f"{file_path.stem}_transcript.txt"
            with open(output_txt, "w", encoding="utf-8") as f:
                f.write(full_transcript)

            end_time = datetime.now()
            processing_time = (end_time - start_time).total_seconds()

            # Create metadata
            metadata = {
                "source_file": file_path.name,
                "transcription_date": end_time.isoformat(),
                "processing_time_seconds": processing_time,
                "model_used": self.model_name,
                "chunk_count": len(chunks),
                "chunk_length_minutes": chunk_length,
                "output_files": {"docx": output_docx.name, "txt": output_txt.name},
            }

            # Save metadata
            metadata_file = output_dir / f"{file_path.stem}_metadata.json"
            self.save_metadata(metadata, metadata_file)

            logger.info(
                f"Processing complete! Time taken: {processing_time:.2f} seconds"
            )
            return metadata

        except Exception as e:
            logger.error(f"Error processing {file_path.name}: {e}")
            raise
        finally:
            # Clean up temp directory
            if temp_dir.exists():
                for temp_file in temp_dir.glob("temp_chunk_*.wav"):
                    temp_file.unlink()
                try:
                    temp_dir.rmdir()
                except OSError:
                    pass  # Directory not empty, leave it

    def process_directory(
        self, input_dir: Path, output_dir: Path, chunk_length: int = 10
    ):
        """Process all audio files in a directory."""
        input_dir = Path(input_dir)
        output_dir = Path(output_dir)

        # Create output directory
        output_dir.mkdir(parents=True, exist_ok=True)

        # Find all supported audio files
        audio_files = []
        for ext in self.supported_formats:
            audio_files.extend(input_dir.glob(f"*{ext}"))

        if not audio_files:
            logger.warning(f"No supported audio files found in {input_dir}")
            return

        logger.info(f"Found {len(audio_files)} audio files to process")

        results = []
        for audio_file in audio_files:
            try:
                metadata = self.process_audio_file(audio_file, output_dir, chunk_length)
                results.append(metadata)
            except Exception as e:
                logger.error(f"Failed to process {audio_file.name}: {e}")
                continue

        # Save batch summary
        if results:
            summary = {
                "batch_date": datetime.now().isoformat(),
                "total_files": len(audio_files),
                "successful_transcriptions": len(results),
                "results": results,
            }
            summary_file = output_dir / "batch_summary.json"
            self.save_metadata(summary, summary_file)


def main():
    parser = argparse.ArgumentParser(
        description="Audio Transcription Tool using Whisper"
    )
    parser.add_argument(
        "--input-dir",
        "-i",
        type=str,
        default=r"C:\Users\Jonathan\Documents\kbdocs\text_result",
        help="Input directory containing audio files",
    )
    parser.add_argument(
        "--output-dir",
        "-o",
        type=str,
        help="Output directory (defaults to input directory)",
    )
    parser.add_argument(
        "--model",
        "-m",
        type=str,
        default="medium.en",
        choices=["tiny.en", "base.en", "small.en", "medium.en", "large"],
        help="Whisper model to use",
    )
    parser.add_argument(
        "--chunk-length", "-c", type=int, default=10, help="Chunk length in minutes"
    )
    parser.add_argument(
        "--device",
        "-d",
        type=str,
        default="cpu",
        choices=["cpu", "cuda"],
        help="Device to use for inference",
    )

    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir) if args.output_dir else input_dir

    if not input_dir.exists():
        logger.error(f"Input directory does not exist: {input_dir}")
        return

    # Initialize transcriber
    transcriber = AudioTranscriber(model_name=args.model, device=args.device)

    # Process directory
    transcriber.process_directory(input_dir, output_dir, args.chunk_length)


if __name__ == "__main__":
    main()
