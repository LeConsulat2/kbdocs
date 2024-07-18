import os
import pypandoc
from unidecode import unidecode
from docx import Document

# Define the directory where the Word files are saved
word_directory = r"C:\Users\Jonathan\Documents\kbdocs\here"

# Define the directory where the PDF files will be saved
pdf_directory = r"C:\Users\Jonathan\Documents\kbdocs\pdfs"

# Check if the PDF directory exists, if not, create it
if not os.path.exists(pdf_directory):
    os.makedirs(pdf_directory)


# Function to clean text using unidecode
def clean_text(text):
    return unidecode(text)


# Iterate through each Word document in the directory
for filename in os.listdir(word_directory):
    if filename.endswith(".docx"):
        # Define the full file path for the Word document
        word_file = os.path.join(word_directory, filename)

        # Define the full file path for the PDF
        pdf_file = os.path.join(pdf_directory, filename.replace(".docx", ".pdf"))

        # Check if the PDF file already exists
        if os.path.exists(pdf_file):
            print(f"PDF already exists for {filename}, skipping conversion.")
            continue  # Skip to the next file

        try:
            # Read the content of the Word document
            doc = Document(word_file)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            text_content = "\n".join(full_text)

            # Use unidecode to handle unsupported characters
            cleaned_content = clean_text(text_content)

            # Write the cleaned content back to a temporary Word document
            temp_word_file = os.path.join(word_directory, "temp.docx")
            temp_doc = Document()
            for paragraph in doc.paragraphs:
                temp_doc.add_paragraph(clean_text(paragraph.text))
            temp_doc.save(temp_word_file)

            # Convert the cleaned Word document to PDF with selectable text
            pypandoc.convert_file(
                temp_word_file,
                "pdf",
                outputfile=pdf_file,
                extra_args=["--pdf-engine=pdflatex"],
            )

            # Remove the temporary file
            os.remove(temp_word_file)

            print(f"Converted {filename} to PDF and saved to {pdf_directory}.")
        except Exception as e:
            print(f"Error converting {filename}: {e}")
            if os.path.exists(temp_word_file):
                os.remove(temp_word_file)

print("All Word documents have been processed.")
