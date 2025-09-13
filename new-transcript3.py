import os
from docx import Document
import whisper
from dotenv import load_dotenv
import openai
from openai import OpenAI
import re
import time

# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Whisper model
print("[INFO] Loading Whisper model (medium.en)...")
whisper_model = whisper.load_model("medium.en", device="cpu")
print("[INFO] Whisper model loaded successfully!")

# Prompt for formatting transcript only (no translation)
FORMATTING_PROMPT = """You are helping to format an English transcript from Nomad Coders Nicolas's lecture into natural subtitle segments.

Your task:
• Take the raw English transcript and break it into natural subtitle chunks (1-2 sentences each)
• Break based on natural speech patterns, emphasis points, and logical pauses
• Don't break in the middle of important concepts or explanations
• Each subtitle segment should be complete and make sense on its own
• Keep the original English text exactly as is - don't modify, correct, or change anything
• Output each subtitle segment on its own line
• Add a blank line between each subtitle segment for spacing

Example input:
"All right everybody in this video we are going to implement our voice agent we're going to turn this agent that is text based into a voice agent now I have to tell you that this implementation of voice agents I think is a little bit rough"

Example output:
All right everybody, in this video we are going to implement our voice agent.

We're going to turn this agent that is text-based into a voice agent.

Now I have to tell you that this implementation of voice agents I think is a little bit rough.

Keep it natural and flowing like subtitles would appear in a video."""


def format_transcript_with_gpt(transcript_text):
    """
    Format transcript into natural subtitle segments using GPT
    """
    try:
        print("[INFO] Formatting transcript into subtitle segments...")

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": FORMATTING_PROMPT},
                {"role": "user", "content": transcript_text},
            ],
            temperature=0.1,  # Very low for consistent formatting
            max_tokens=15000,
        )

        formatted_transcript = response.choices[0].message.content
        print("[INFO] Transcript formatting complete.")
        return formatted_transcript

    except Exception as e:
        print(f"[ERROR] Formatting failed: {e}")
        # Fallback: basic sentence splitting
        sentences = transcript_text.split(".")
        formatted = ""
        for i, sentence in enumerate(sentences):
            if sentence.strip():
                formatted += sentence.strip() + "."
                if i < len(sentences) - 1:
                    formatted += "\n\n"
        return formatted


def transcribe_and_format_audio(file_path, output_directory):
    """
    Main function that transcribes and formats transcript only
    """
    # Transcribe audio
    print(f"[INFO] Transcribing '{file_path}'...")
    result = whisper_model.transcribe(file_path, word_timestamps=True)

    # Get full transcript text
    full_transcript = result["text"].strip()

    if not full_transcript:
        print("[WARNING] No transcription found.")
        return ""

    print(f"[INFO] Transcription length: {len(full_transcript)} characters")

    # Format the transcript into subtitle segments
    formatted_content = format_transcript_with_gpt(full_transcript)

    return formatted_content


def save_transcript_to_docx(content, output_file):
    """
    Save formatted transcript to DOCX
    """
    print(f"[INFO] Saving transcript to '{output_file}'...")
    doc = Document()

    # Split content by lines and handle spacing
    lines = content.split("\n")

    for line in lines:
        if line.strip():
            # Add paragraph with content
            doc.add_paragraph(line.strip())
        else:
            # Add empty paragraph for spacing
            doc.add_paragraph()

    doc.save(output_file)
    print(f"[INFO] Transcript saved successfully!")


def save_transcript_to_txt(content, output_file):
    """
    Save formatted transcript to TXT file as well
    """
    print(f"[INFO] Saving transcript to '{output_file}'...")
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"[INFO] TXT file saved successfully!")


# Main execution
if __name__ == "__main__":
    directory = r"C:\Users\Jonathan\Documents\kbdocs\text_result"
    media_files = [f for f in os.listdir(directory) if f.endswith((".mp4", ".mp3"))]

    if not media_files:
        print("[WARNING] No media files found in the directory.")
    else:
        for media_file in media_files:
            audio_file_path = os.path.join(directory, media_file)
            print(f"[INFO] Processing file: {audio_file_path}")

            # Get formatted transcript
            formatted_transcript = transcribe_and_format_audio(
                audio_file_path, directory
            )

            if formatted_transcript:
                print("[INFO] Transcription and formatting completed!")

                # Save to DOCX
                base_name = os.path.splitext(media_file)[0]
                docx_output = os.path.join(directory, f"{base_name}_transcript.docx")
                save_transcript_to_docx(formatted_transcript, docx_output)

                # Also save to TXT for easy reading
                # txt_output = os.path.join(directory, f"{base_name}_transcript.txt")
                # save_transcript_to_txt(formatted_transcript, txt_output)

                print(f"[INFO] Files saved:")
                print(f"  - DOCX: {docx_output}")
                # print(f"  - TXT: {txt_output}")
            else:
                print("[WARNING] No transcript content to save.")
