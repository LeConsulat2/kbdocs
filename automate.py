import os
import re
from pathlib import Path
from typing import List, Tuple
import logging
from datetime import datetime

from pydub import AudioSegment
from openai import OpenAI
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('subtitle_processing.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

class SubtitleStyleProcessor:
    def __init__(self, transcribe_model: str = "gpt-4o-transcribe", translate_model: str = "gpt-4o"):
        self.transcribe_model = transcribe_model
        self.translate_model = translate_model
        self.max_file_size_mb = 25
        
        # Nomad Coders specific patterns
        self.short_phrases = [
            "okay", "alright", "fantastic", "perfect", "cool", "great", "nice", 
            "good", "yes", "yep", "no", "nope", "right", "exactly", "boom",
            "let's go", "here we go", "there we go", "all right", "very good"
        ]
        
        self.connectors = [
            "so", "now", "and", "but", "because", "if", "when", "as you can see",
            "for example", "like this", "let's", "we're going to", "I will",
            "you can", "we can", "this is", "that is"
        ]

    def transcribe_with_timestamps(self, audio_file_path: str) -> dict:
        """Transcribe with detailed timestamps for better segmentation."""
        try:
            with open(audio_file_path, "rb") as audio_file:
                response = client.audio.transcriptions.create(
                    model=self.transcribe_model,
                    file=audio_file,
                    language="en",
                    prompt="This is Nomad Coders coding lecture by Nicolas. He speaks in short phrases with pauses. Please maintain natural speech boundaries and include filler words like 'okay', 'alright', 'fantastic'.",
                    response_format="verbose_json",
                    temperature=0.0,
                    timestamp_granularities=["word"]
                )
                
                return response
                
        except Exception as e:
            logger.error(f"Error transcribing with timestamps: {e}")
            # Fallback without timestamps
            with open(audio_file_path, "rb") as audio_file:
                response = client.audio.transcriptions.create(
                    model=self.transcribe_model,
                    file=audio_file,
                    language="en",
                    response_format="text",
                    temperature=0.0
                )
                return {"text": response}

    def smart_sentence_split(self, text: str, words_data: List = None) -> List[str]:
        """
        Smart splitting that mimics Nicolas's natural speech patterns.
        Tries to create subtitle-length segments.
        """
        # First, clean the text
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Define patterns for natural breaks
        natural_breaks = [
            # Strong breaks
            r'(?<=[.!?])\s+(?=[A-Z])',  # Sentence endings
            r'(?<=\w)\s+(?=okay|alright|fantastic|perfect|cool|great|nice|good)\b',  # Before affirmations
            r'\b(?:okay|alright|fantastic|perfect|cool|great|nice|good)\s+(?=\w)',  # After affirmations
            
            # Medium breaks  
            r'\b(?:so|now|and|but|because|if|when)\s+',  # Connectors
            r'\b(?:as you can see|for example|like this|let\'s|we\'re going to|I will)\s+',  # Common phrases
            
            # Weak breaks (only if segment is getting too long)
            r'(?<=\w)\s+(?=\w)',  # Any word boundary
        ]
        
        segments = [text]
        
        # Apply breaks in order of strength
        for pattern in natural_breaks:
            new_segments = []
            for segment in segments:
                if len(segment) > 80:  # Only split if segment is long
                    parts = re.split(f'({pattern})', segment)
                    current = ""
                    for part in parts:
                        if current and len(current + part) > 120:  # Max subtitle length
                            if current.strip():
                                new_segments.append(current.strip())
                            current = part
                        else:
                            current += part
                    if current.strip():
                        new_segments.append(current.strip())
                else:
                    new_segments.append(segment)
            segments = new_segments
        
        # Clean up segments
        final_segments = []
        for segment in segments:
            segment = segment.strip()
            if segment and len(segment) > 3:  # Filter out very short meaningless segments
                final_segments.append(segment)
        
        return final_segments

    def translate_segment_precise(self, english_segment: str) -> str:
        """
        Translate a single segment with precise Nomad Coders style.
        """
        translation_prompt = f"""너는 노마드코더 니콜라스의 강의 번역 전문가야. 

**중요 규칙:**
1. 입력받은 영어 문장을 절대 수정하지 말고 그대로 유지해
2. 니콜라스의 캐주얼하고 친근한 말투로 번역해
3. 기술 용어는 번역하지 않고 원문 그대로 써
4. 끝맺음: "해, 할 거야, 하는 거지, 그렇고" 등 자연스럽게
5. "사용자" → "유저", "ID" → "_id" 형태로

**출력 형식:**
**원문 그대로**
번역된 한국어

다음을 번역해:
{english_segment}"""

        try:
            response = client.chat.completions.create(
                model=self.translate_model,
                messages=[
                    {"role": "user", "content": translation_prompt}
                ],
                temperature=0.1,
                max_tokens=300
            )
            
            result = response.choices[0].message.content.strip()
            
            # Parse the result to extract English and Korean parts
            lines = result.split('\n')
            
            # Look for the pattern: **English** followed by Korean
            english_line = ""
            korean_line = ""
            
            for i, line in enumerate(lines):
                line = line.strip()
                if line.startswith('**') and line.endswith('**'):
                    english_line = line
                    # Next non-empty line should be Korean
                    for j in range(i+1, len(lines)):
                        if lines[j].strip():
                            korean_line = lines[j].strip()
                            break
                    break
            
            # If parsing failed, create format manually
            if not english_line:
                english_line = f"**{english_segment}**"
            if not korean_line:
                korean_line = "번역 처리 중 오류"
            
            return f"{english_line}\n{korean_line}"
            
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return f"**{english_segment}**\n번역 오류 발생"

    def process_transcription_to_subtitles(self, transcription_data: dict) -> str:
        """
        Convert transcription to subtitle format.
        """
        if isinstance(transcription_data, dict) and 'text' in transcription_data:
            text = transcription_data['text']
            words_data = transcription_data.get('words', [])
        else:
            text = str(transcription_data)
            words_data = []
        
        logger.info("Splitting transcription into subtitle segments...")
        segments = self.smart_sentence_split(text, words_data)
        
        logger.info(f"Created {len(segments)} segments. Starting translation...")
        
        subtitle_content = ""
        
        for i, segment in enumerate(segments):
            logger.info(f"Translating segment {i+1}/{len(segments)}: {segment[:50]}...")
            
            translated_pair = self.translate_segment_precise(segment)
            subtitle_content += translated_pair + "\n"
            
            # Add spacing between segments for readability
            if i < len(segments) - 1:
                subtitle_content += "\n"
        
        return subtitle_content.strip()

    def convert_to_wav(self, file_path: str, output_directory: str) -> str:
        """Convert audio file to WAV format."""
        logger.info(f"Converting '{file_path}' to WAV format...")
        
        try:
            audio = AudioSegment.from_file(file_path)
            audio = audio.set_frame_rate(16000).set_channels(1)
            
            wav_file_path = os.path.join(
                output_directory, 
                Path(file_path).stem + "_temp.wav"
            )
            audio.export(wav_file_path, format="wav")
            logger.info(f"Conversion complete: '{wav_file_path}'")
            return wav_file_path
        except Exception as e:
            logger.error(f"Error converting file: {e}")
            raise

    def save_subtitle_files(self, content: str, base_path: str, original_filename: str):
        """Save subtitle content in multiple formats."""
        
        # Save as TXT (main format)
        txt_file = f"{base_path}_subtitles.txt"
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Subtitles saved: {txt_file}")
        
        # Save as DOCX with formatting
        docx_file = f"{base_path}_subtitles.docx"
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Nomad Coders Subtitles: {Path(original_filename).stem}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        metadata = doc.add_paragraph()
        metadata.add_run('Generated: ').bold = True
        metadata.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        metadata.add_run('\nOriginal: ').bold = True
        metadata.add_run(original_filename)
        
        doc.add_page_break()
        
        # Process content
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('**') and line.endswith('**'):
                # English line (bold)
                para = doc.add_paragraph()
                para.add_run(line[2:-2]).bold = True
            else:
                # Korean line (normal)
                para = doc.add_paragraph(line)

        doc.save(docx_file)
        logger.info(f"Formatted document saved: {docx_file}")

    def process_file(self, file_path: str, output_directory: str):
        """Main processing function."""
        logger.info(f"Starting subtitle processing: {file_path}")
        
        # Create temp directory
        temp_dir = os.path.join(output_directory, "temp")
        os.makedirs(temp_dir, exist_ok=True)
        
        try:
            # Convert to WAV if needed
            if not file_path.lower().endswith('.wav'):
                wav_file = self.convert_to_wav(file_path, temp_dir)
            else:
                wav_file = file_path
            
            # Check file size
            file_size_mb = os.path.getsize(wav_file) / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                logger.warning(f"File too large ({file_size_mb:.1f}MB). Consider splitting manually.")
                # You could add chunking logic here if needed
            
            # Transcribe with timestamps
            logger.info("Transcribing audio...")
            transcription_data = self.transcribe_with_timestamps(wav_file)
            
            # Process to subtitle format
            subtitle_content = self.process_transcription_to_subtitles(transcription_data)
            
            # Save results
            base_name = Path(file_path).stem
            base_path = os.path.join(output_directory, base_name)
            
            self.save_subtitle_files(subtitle_content, base_path, os.path.basename(file_path))
            
            logger.info(f"Subtitle processing complete for: {file_path}")
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            raise
        finally:
            # Cleanup temp files
            if os.path.exists(temp_dir):
                for temp_file in os.listdir(temp_dir):
                    try:
                        os.remove(os.path.join(temp_dir, temp_file))
                    except:
                        pass
                try:
                    os.rmdir(temp_dir)
                except:
                    pass

def main():
    """Main execution."""
    # Configuration
    directory = r"C:\Users\Jonathan\Documents\kbdocs\text_result"
    supported_formats = ('.mp4', '.mp3', '.wav', '.m4a')
    
    # Initialize processor
    processor = SubtitleStyleProcessor(
        transcribe_model="gpt-4o-transcribe",
        translate_model="gpt-4o"
    )
    
    # Find files
    media_files = [
        f for f in os.listdir(directory) 
        if f.lower().endswith(supported_formats)
    ]
    
    if not media_files:
        logger.warning("No media files found.")
        return
    
    logger.info(f"Found {len(media_files)} files to process")
    
    # Process each file
    for media_file in media_files:
        file_path = os.path.join(directory, media_file)
        try:
            logger.info(f"Processing: {media_file}")
            processor.process_file(file_path, directory)
            logger.info(f"✓ Completed: {media_file}")
        except Exception as e:
            logger.error(f"✗ Failed {media_file}: {e}")
            continue
    
    logger.info("All processing complete!")

if __name__ == "__main__":
    main()