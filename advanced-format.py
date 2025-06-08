import os
import re
from pathlib import Path
from typing import List, Tuple
import logging
from datetime import datetime

from pydub import AudioSegment
from openai import OpenAI
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv

# --- 추가된 임포트 ---
import whisper
# --------------------

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('subtitle_processing.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Initialize OpenAI client (for translation only)
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

class SubtitleStyleProcessor:
    # --- transcribe_model 대신 local_whisper_model_name 추가 ---
    def __init__(self, local_whisper_model_name: str = "medium", translate_model: str = "gpt-4o"):
        self.transcribe_model_name = local_whisper_model_name # 모델 이름만 저장
        self.translate_model = translate_model
        self.max_file_size_mb = 25 # 이 값은 OpenAI API 사용 시 유효하며, 로컬 Whisper는 메모리 허용 범위 내에서 더 큰 파일도 처리 가능

        logger.info(f"Loading local Whisper model: '{self.transcribe_model_name}'...")
        try:
            self.whisper_model = whisper.load_model(self.transcribe_model_name)
            logger.info(f"Local Whisper model '{self.transcribe_model_name}' loaded successfully.")
        except Exception as e:
            logger.error(f"Error loading local Whisper model: {e}")
            logger.error("Please ensure 'openai-whisper' is installed (`pip install openai-whisper`) "
                         "and ffmpeg is in your system's PATH.")
            raise # 모델 로드 실패 시 종료

        # Nomad Coders specific patterns (변동 없음)
        self.short_phrases = [
            "okay", "alright", "fantastic", "perfect", "cool", "great", "nice", 
            "good", "yes", "yep", "no", "nope", "right", "exactly", "boom",
            "let's go", "here we go", "there we go", "all right", "very good"
        ]
        
        self.connectors = [
            "so", "now", "and", "but", "because", "if", "when", "as you can see",
            "for example", "like this", "let's", "we're going to", "I will",
            "you can", "we can", "this is", "that is"
        ]

    def transcribe_with_timestamps(self, audio_file_path: str) -> dict:
        """
        Transcribe with detailed timestamps using the local Whisper model.
        Returns a dictionary mimicking OpenAI's verbose_json for compatibility.
        """
        logger.info(f"Transcribing audio with local Whisper model '{self.transcribe_model_name}'...")
        try:
            # Local Whisper 모델을 사용하여 전사
            # word_timestamps=True는 단어별 타임스탬프를 제공합니다.
            # initial_prompt는 OpenAI의 'prompt'와 유사하게 동작합니다.
            result = self.whisper_model.transcribe(
                audio_file_path,
                word_timestamps=True,
                language="en",
                initial_prompt="This is Nomad Coders coding lecture by Nicolas. He speaks in short phrases with pauses. Please maintain natural speech boundaries and include filler words like 'okay', 'alright', 'fantastic'.",
                verbose=False # Whisper 자체의 상세 출력을 억제
            )
            
            # 로컬 Whisper의 출력 형식을 OpenAI verbose_json과 유사하게 변환
            # OpenAI verbose_json은 'words' 키에 모든 단어 리스트를 포함합니다.
            # Local Whisper는 'segments' 안에 각 세그먼트의 'words'를 포함합니다.
            
            all_words = []
            full_text = ""
            if 'segments' in result:
                for segment in result['segments']:
                    full_text += segment['text'] + " "
                    if 'words' in segment:
                        all_words.extend(segment['words'])
            else: # Fallback if 'segments' is not present (unlikely with word_timestamps=True)
                full_text = result.get('text', '')

            # 'text'와 'words' 키를 가진 딕셔너리로 반환하여 `process_transcription_to_subtitles`와 호환되게 함
            return {"text": full_text.strip(), "words": all_words}
                
        except Exception as e:
            logger.error(f"Error transcribing with local Whisper: {e}")
            # 로컬 Whisper 사용 시에는 파일 오픈 에러 외에는 바로 텍스트로 폴백하는 경우가 적음
            # 대신 에러 발생 시 빈 텍스트라도 반환하여 다음 단계가 진행될 수 있도록 처리
            return {"text": "", "words": []}

    def smart_sentence_split(self, text: str, words_data: List = None) -> List[str]:
        """
        Smart splitting that mimics Nicolas's natural speech patterns.
        Tries to create subtitle-length segments.
        Note: This function currently uses text-based splitting.
        To use 'words_data' (timestamps) for more accurate, speech-timed splitting,
        this function would need a significant rewrite.
        """
        # First, clean the text
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Define patterns for natural breaks
        natural_breaks = [
            # Strong breaks
            r'(?<=[.!?])\s+(?=[A-Z])',  # Sentence endings
            r'(?<=\w)\s+(?=okay|alright|fantastic|perfect|cool|great|nice|good)\b',  # Before affirmations
            r'\b(?:okay|alright|fantastic|perfect|cool|great|nice|good)\s+(?=\w)',  # After affirmations
            
            # Medium breaks  
            r'\b(?:so|now|and|but|because|if|when)\s+',  # Connectors
            r'\b(?:as you can see|for example|like this|let\'s|we\'re going to|I will)\s+',  # Common phrases
            
            # Weak breaks (only if segment is getting too long)
            r'(?<=\w)\s+(?=\w)',  # Any word boundary
        ]
        
        segments = [text]
        
        # Apply breaks in order of strength
        for pattern in natural_breaks:
            new_segments = []
            for segment in segments:
                if len(segment) > 80:  # Only split if segment is long
                    parts = re.split(f'({pattern})', segment)
                    current = ""
                    for part in parts:
                        if current and len(current + part) > 120:  # Max subtitle length
                            if current.strip():
                                new_segments.append(current.strip())
                            current = part
                        else:
                            current += part
                    if current.strip():
                        new_segments.append(current.strip())
                else:
                    new_segments.append(segment)
            segments = new_segments
        
        # Clean up segments
        final_segments = []
        for segment in segments:
            segment = segment.strip()
            if segment and len(segment) > 3:  # Filter out very short meaningless segments
                final_segments.append(segment)
        
        return final_segments

    def translate_segment_precise(self, english_segment: str) -> str:
        """
        Translate a single segment with precise Nomad Coders style.
        (이 함수는 OpenAI API를 그대로 사용합니다)
        """
        translation_prompt = f"""너는 노마드코더(Nomad Coders)의 니콜라스 강의를 번역하는 역할이야.
니콜라스의 스타일은 다음과 같아:
•	캐주얼하고 재미있는 말투
•	실용적이고 직관적인 설명
•	격식 없는 자연스러운 대화체
번역 규칙
•	입력된 영어 문장(라인)은 절대 수정하지 않고 그대로 둬.
•	영어 문장은 굵게 표시 (bold), 번역된 한국어 문장은 평범하게 표시.
•	출력은 영어 문장 바로 아래에 한국어 번역이 나와야 해.
•	번역은 자연스럽고 캐주얼하게 진행해. 끝맺음은 "해, 할 거야, 하는 거지, 그렇고"처럼 대화하듯이 작성해.
•	헤딩, 구분선 같은 건 넣지 말고 깔끔하게 이어서 작성해.
•	전문 용어는 번역하지 않고 원문 그대로 작성해. 
o	예: console log, deploy, import, variable, function, constraints, type, Typescript, foreign key, primary key, column, composite 등 리스트에 없는 경우도 문맥을 보고 니콜라스 스타일에 맞게 판단해. 특정한 행동이나 작성 지시가 포함된 경우 예: "위 같이 작성," "화면 참조," "추가사항 적용" 등의 표현은 문맥에 따라 간결하게 자연스럽게 설명해. '사용자'라는 표현 대신 '유저'를 사용해. 
•	Profile ID, User ID , Product ID 등등 그 이 외에도 이런식의 문맥이 나오면 profile_id, user_id, product_id 등등 이런식으로 대답해
출력 예시
In this video I want to share with you what are the ingredients we are going to use to build our startup project.
이번 영상에서는 우리가 스타트업 프로젝트를 만들 때 쓸 핵심 요소들을 알려줄 거야.
Now originally this was a course based on Remix.
원래 이 강의는 Remix를 기반으로 만든 거야.
This table references the post ID and the profile ID. 
이 테이블은 post_id와 profile_id를 참조해.


다음을 번역해:
{english_segment}"""

        try:
            response = client.chat.completions.create(
                model=self.translate_model,
                messages=[
                    {"role": "user", "content": translation_prompt}
                ],
                temperature=0.1,
                max_tokens=300
            )
            
            result = response.choices[0].message.content.strip()
            
            # Parse the result to extract English and Korean parts
            lines = result.split('\n')
            
            # Look for the pattern: **English** followed by Korean
            english_line = ""
            korean_line = ""
            
            for i, line in enumerate(lines):
                line = line.strip()
                if line.startswith('**') and line.endswith('**'):
                    english_line = line
                    # Next non-empty line should be Korean
                    for j in range(i+1, len(lines)):
                        if lines[j].strip():
                            korean_line = lines[j].strip()
                            break
                    break
            
            # If parsing failed, create format manually
            if not english_line:
                english_line = f"**{english_segment}**"
            if not korean_line:
                korean_line = "번역 처리 중 오류"
            
            return f"{english_line}\n{korean_line}"
            
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return f"**{english_segment}**\n번역 오류 발생"

    def process_transcription_to_subtitles(self, transcription_data: dict) -> str:
        """
        Convert transcription to subtitle format.
        """
        if isinstance(transcription_data, dict) and 'text' in transcription_data:
            text = transcription_data['text']
            words_data = transcription_data.get('words', [])
        else:
            text = str(transcription_data)
            words_data = []
        
        logger.info("Splitting transcription into subtitle segments...")
        segments = self.smart_sentence_split(text, words_data)
        
        logger.info(f"Created {len(segments)} segments. Starting translation...")
        
        subtitle_content = ""
        
        for i, segment in enumerate(segments):
            logger.info(f"Translating segment {i+1}/{len(segments)}: {segment[:50]}...")
            
            translated_pair = self.translate_segment_precise(segment)
            subtitle_content += translated_pair + "\n"
            
            # Add spacing between segments for readability
            if i < len(segments) - 1:
                subtitle_content += "\n"
        
        return subtitle_content.strip()

    def convert_to_wav(self, file_path: str, output_directory: str) -> str:
        """Convert audio file to WAV format."""
        logger.info(f"Converting '{file_path}' to WAV format...")
        
        try:
            audio = AudioSegment.from_file(file_path)
            # Whisper는 16kHz, 모노 채널을 선호하며, 자동으로 변환하긴 하지만 명시적으로 설정하는 것이 좋습니다.
            audio = audio.set_frame_rate(16000).set_channels(1) 
            
            wav_file_path = os.path.join(
                output_directory, 
                Path(file_path).stem + "_temp.wav"
            )
            audio.export(wav_file_path, format="wav")
            logger.info(f"Conversion complete: '{wav_file_path}'")
            return wav_file_path
        except Exception as e:
            logger.error(f"Error converting file: {e}")
            raise

    def save_subtitle_files(self, content: str, base_path: str, original_filename: str):
        """Save subtitle content in multiple formats."""
        
        # Save as TXT (main format)
        txt_file = f"{base_path}_subtitles.txt"
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Subtitles saved: {txt_file}")
        
        # Save as DOCX with formatting
        docx_file = f"{base_path}_subtitles.docx"
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Nomad Coders Subtitles: {Path(original_filename).stem}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        metadata = doc.add_paragraph()
        metadata.add_run('Generated: ').bold = True
        metadata.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        metadata.add_run('\nOriginal: ').bold = True
        metadata.add_run(original_filename)
        
        doc.add_page_break()
        
        # Process content
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('**') and line.endswith('**'):
                # English line (bold)
                para = doc.add_paragraph()
                para.add_run(line[2:-2]).bold = True
            else:
                # Korean line (normal)
                para = doc.add_paragraph(line)

        doc.save(docx_file)
        logger.info(f"Formatted document saved: {docx_file}")

    def process_file(self, file_path: str, output_directory: str):
        """Main processing function."""
        logger.info(f"Starting subtitle processing: {file_path}")
        
        # Create temp directory
        temp_dir = os.path.join(output_directory, "temp")
        os.makedirs(temp_dir, exist_ok=True)
        
        try:
            # Convert to WAV if needed
            # 로컬 Whisper는 다양한 오디오 포맷을 지원하지만,
            # pydub를 통한 WAV 변환은 안정성과 일관성을 위해 유지합니다.
            if not file_path.lower().endswith('.wav'):
                wav_file = self.convert_to_wav(file_path, temp_dir)
            else:
                wav_file = file_path
            
            # Check file size (로컬 Whisper는 25MB 제한을 받지 않으나, 메모리 사용량은 고려해야 함)
            file_size_mb = os.path.getsize(wav_file) / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                logger.warning(f"File size is {file_size_mb:.1f}MB, which is large. "
                               "Local Whisper can handle larger files, but performance depends on hardware. "
                               "Consider manually splitting very large files if issues occur.")
                # 긴 파일 처리 시 여기서 오디오 청킹 로직을 추가할 수 있습니다.
            
            # Transcribe with timestamps
            transcription_data = self.transcribe_with_timestamps(wav_file)
            
            # Process to subtitle format
            subtitle_content = self.process_transcription_to_subtitles(transcription_data)
            
            # Save results
            base_name = Path(file_path).stem
            base_path = os.path.join(output_directory, base_name)
            
            self.save_subtitle_files(subtitle_content, base_path, os.path.basename(file_path))
            
            logger.info(f"Subtitle processing complete for: {file_path}")
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            raise
        finally:
            # Cleanup temp files
            if os.path.exists(temp_dir):
                for temp_file in os.listdir(temp_dir):
                    try:
                        os.remove(os.path.join(temp_dir, temp_file))
                    except:
                        pass
                try:
                    os.rmdir(temp_dir)
                except:
                    pass

def main():
    """Main execution."""
    # Configuration
    directory = r"C:\Users\Jonathan\Documents\kbdocs\text_result"
    supported_formats = ('.mp4', '.mp3', '.wav', '.m4a')
    
    # Initialize processor
    # local_whisper_model_name: "tiny", "base", "small", "medium", "large", "large-v1", "large-v2", "large-v3"
    # 모델 크기가 커질수록 정확도는 높아지지만, 필요한 메모리(VRAM)와 처리 시간도 증가합니다.
    # GPU가 없다면 "base" 또는 "small" 모델이 적절할 수 있습니다.
    processor = SubtitleStyleProcessor(
        local_whisper_model_name="base", # 여기에 사용할 로컬 Whisper 모델 이름 지정
        translate_model="gpt-4o"
    )
    
    # Find files
    media_files = [
        f for f in os.listdir(directory) 
        if f.lower().endswith(supported_formats)
    ]
    
    if not media_files:
        logger.warning("No media files found.")
        return
    
    logger.info(f"Found {len(media_files)} files to process")
    
    # Process each file
    for media_file in media_files:
        file_path = os.path.join(directory, media_file)
        try:
            logger.info(f"Processing: {media_file}")
            processor.process_file(file_path, directory)
            logger.info(f"✓ Completed: {media_file}")
        except Exception as e:
            logger.error(f"✗ Failed {media_file}: {e}")
            continue
    
    logger.info("All processing complete!")

if __name__ == "__main__":
    main()