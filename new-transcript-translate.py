import os
from docx import Document
import whisper
from dotenv import load_dotenv
import openai
from openai import OpenAI
import re
import time

# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Whisper model
print("[INFO] Loading Whisper model (medium.en)...")
whisper_model = whisper.load_model("medium.en", device="cpu")
print("[INFO] Whisper model loaded successfully!")

# Updated translation prompt with user's improvements
TRANSLATION_PROMPT = """너는 노마드코더(Nomad Coders)의 니콜라스 강의를 번역하는 역할이야.

니콜라스의 스타일은 다음과 같아:
• 캐주얼하고 재미있는 말투
• 실용적이고 직관적인 설명
• 격식 없는 자연스러운 대화체

번역 규칙:
• 입력된 transcript(영어 원문)는 한 번에 여러 줄로 주어질 수 있어. 네 역할은 이를 **자연스러운 자막 단위(짧고 간결한 구간)**로 끊어서 정리하는 거야.
• 한 문장을 무조건 하나로 자르지 말고, 맥락이나 강조 포인트 단위로 1~2줄씩 끊어줘.
• 영어 문장은 절대 수정하지 않고 그대로 두고, 굵게 표시 (bold).
• 바로 아래에 한국어 번역을 붙여. 한국어 번역은 자연스럽고 캐주얼하게, "해, 할 거야, 하는 거지, 그렇고" 같은 대화체로 작성해.
• 헤딩, 구분선 같은 건 넣지 말고 깔끔하게 이어서 작성해.
• you라는 단어 나오면 "여러분", "여러분이", 아니면 문맥상 자연스러우면 여러분 넣지 않고 번역해.
• 전문 용어(import, function, class, async/await, coroutine, decorator, variable, API, endpoint, JSON, YAML, pip, venv, package manager, CrewAI, AutoGen, OpenAI Agents SDK, LangGraph, Google Agent Builder, kit, LLM, prompt, system message, user message, state, tool, orchestrator 등)는 번역하지 않고 그대로 둬.
• snake_case / camelCase / kebab-case / SNAKE_CASE 등은 그대로 유지 (예: transfer_agent, web_search_tool, job_search_agent).
• agent라고만 딱 나오면 한국말로 "에이전트"로 쓰고, 다른 전문용어는 그대로 써.
• 코드 블록(python, bash, json 등)은 번역하지 않고 그대로 둬.
• 특정한 행동이나 작성 지시가 포함된 경우 (예: "위 같이 작성", "화면 참조", "추가사항 적용")는 문맥에 맞게 간결하고 자연스럽게 설명해.
• "사용자"라는 표현 대신 항상 "유저"라고 번역해.

중요 (VERY IMPORTANT):
• 너무 흔한 단어들(state, function calling, import, class 등)은 그대로 두고 (한국말) 붙이지 않아.
• 강의에서 처음 나오거나 헷갈릴 수 있는 개념은 → 처음 등장할 때만 짧고 직관적인 한국어 설명을 () 안에 넣어줘.
  예:
  - Conditional_Edge(조건에 따라 다음 Node로 이동 여부를 결정하는 엣지)
  - Edge(엣지, 노드 간 연결선)
  - Node(노드, 작업이나 상태를 담는 단위 박스)
  - Decorator(데코레이터, 함수/클래스에 기능을 쉽게 덧붙이는 문법 도구)
  이후 다시 나올 땐 영어만 써.

출력 예시:
**All right everybody it is now time to set up our environment**
좋아 모두, 이제 environment(환경)를 설정할 시간이야.

**and not to set up one environment**
근데 딱 하나만 설정하는 건 아니야.

**I just want to show you how to set up a environment**
내가 보여주려는 건 environment 만드는 방법이야.

**because the way we're going to work in this course is by dividing and conquering okay**
왜냐면 이번 코스는 divide and conquer(나눠서 정복하기) 방식으로 진행할 거거든.
"""


def split_by_sentences_natural(transcription_segments, max_duration=45):
    """
    Split transcription into natural subtitle chunks with longer duration
    """
    chunks = []
    current_chunk = {"text": "", "start": None, "end": None}

    for segment in transcription_segments:
        # If adding this segment would make chunk too long, finish current chunk
        if (
            current_chunk["start"] is not None
            and segment["end"] - current_chunk["start"] > max_duration
        ):
            if current_chunk["text"].strip():
                chunks.append(current_chunk)

            current_chunk = {
                "text": segment["text"].strip(),
                "start": segment["start"],
                "end": segment["end"],
            }
        else:
            # Add to current chunk
            if current_chunk["start"] is None:
                current_chunk["start"] = segment["start"]
            current_chunk["end"] = segment["end"]
            current_chunk["text"] += " " + segment["text"].strip()

    # Add final chunk
    if current_chunk["text"].strip():
        chunks.append(current_chunk)

    return chunks


def translate_full_text_with_gpt(full_text):
    """
    Translate the entire text at once for better flow and consistency
    """
    try:
        print("[INFO] Translating full transcription...")

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": TRANSLATION_PROMPT},
                {"role": "user", "content": full_text},
            ],
            temperature=0.2,  # Lower temperature for consistency
            max_tokens=20000,
        )

        translation = response.choices[0].message.content
        print("[INFO] Full translation complete.")
        return translation

    except Exception as e:
        print(f"[ERROR] Translation failed: {e}")
        # Fallback: return original text with basic formatting
        lines = full_text.split(".")
        formatted = ""
        for line in lines:
            if line.strip():
                formatted += f"**{line.strip()}**\n번역 실패\n\n"
        return formatted


def transcribe_and_translate_audio(file_path, output_directory, chunk_duration=45):
    """
    Main function that transcribes and translates with improved flow
    """
    # Transcribe directly with word-level timestamps
    print(f"[INFO] Transcribing '{file_path}' directly with timestamps...")
    result = whisper_model.transcribe(file_path, word_timestamps=True)

    # Get full text for better translation
    full_text = result["text"].strip()

    if not full_text:
        print("[WARNING] No transcription found.")
        return ""

    print(f"[INFO] Full transcription length: {len(full_text)} characters")

    # Translate the entire text at once for better flow
    translated_content = translate_full_text_with_gpt(full_text)

    return translated_content


def save_to_docx_improved(content, output_file):
    """
    Improved DOCX saving with better formatting
    """
    print(f"[INFO] Saving to '{output_file}'...")
    doc = Document()

    # Split content by lines
    lines = content.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            # Add empty paragraph for spacing
            doc.add_paragraph()
            continue

        # Check if line starts with ** (English text)
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            # Bold English text
            english_text = line[2:-2]  # Remove ** from both ends
            p = doc.add_paragraph()
            run = p.add_run(english_text)
            run.bold = True
        else:
            # Regular Korean translation
            doc.add_paragraph(line)

    doc.save(output_file)
    print(f"[INFO] File saved successfully!")


# Main execution
if __name__ == "__main__":
    directory = r"C:\Users\Jonathan\Documents\kbdocs\text_result"
    media_files = [f for f in os.listdir(directory) if f.endswith((".mp4", ".mp3"))]

    if not media_files:
        print("[WARNING] No media files found in the directory.")
    else:
        for media_file in media_files:
            audio_file_path = os.path.join(directory, media_file)
            print(f"[INFO] Processing file: {audio_file_path}")

            # Process with longer chunks for better translation flow
            translated_content = transcribe_and_translate_audio(
                audio_file_path, directory, chunk_duration=60
            )

            if translated_content:
                print("[INFO] Transcription and translation completed!")

                # Save to file
                output_file_name = os.path.splitext(media_file)[0] + "_translated.docx"
                output_file_path = os.path.join(directory, output_file_name)
                save_to_docx_improved(translated_content, output_file_path)

                print(f"[INFO] Final output saved to: {output_file_path}")
            else:
                print("[WARNING] No content to save.")
