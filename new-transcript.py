import os
from docx import Document
import whisper
from dotenv import load_dotenv
import openai
from openai import OpenAI
import re
import time

# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Whisper model
print("[INFO] Loading Whisper model (medium.en)...")
whisper_model = whisper.load_model("medium.en", device="cpu")
print("[INFO] Whisper model loaded successfully!")

# Translation prompt
TRANSLATION_PROMPT = """너는 노마드코더(Nomad Coders)의 니콜라스 강의를 번역하는 역할이야.

니콜라스의 스타일은 다음과 같아:
• 캐주얼하고 재미있는 말투
• 실용적이고 직관적인 설명
• 격식 없는 자연스러운 대화체

번역 규칙
• 입력된 transcript(영어 원문)는 한 번에 여러 줄로 주어질 수 있어. 네 역할은 이를 **자연스러운 자막 단위(짧고 간결한 구간)**로 끊어서 정리하는 거야.
• 한 문장을 무조건 하나로 자르지 말고, 맥락이나 강조 포인트 단위로 1~2줄씩 끊어줘.
• 영어 문장은 절대 수정하지 않고 그대로 두고, 굵게 표시 (bold).
• 바로 아래에 한국어 번역을 붙여. 한국어 번역은 자연스럽고 캐주얼하게, "해, 할 거야, 하는 거지, 그렇고" 같은 대화체로 작성해.
• 헤딩, 구분선 같은 건 넣지 말고 깔끔하게 이어서 작성해.
• 전문 용어(import, function, class, async/await, coroutine, decorator, variable, API, endpoint, JSON, YAML, pip, venv, package manager, CrewAI, AutoGen, OpenAI Agents SDK, LangGraph, Google Agent Builder, kit, LLM, prompt, system message, user message, state, tool, agent, workflow, orchestrator 등)는 번역하지 않고 그대로 둬.
• snake_case / camelCase / kebab-case / SNAKE_CASE 등은 그대로 유지 (예: profile_id, user_id, product_id).
• 코드 블록(python, bash, json 등)은 번역하지 않고 그대로 둬.
• 특정한 행동이나 작성 지시가 포함된 경우 (예: "위 같이 작성," "화면 참조," "추가사항 적용")는 문맥에 맞게 간결하고 자연스럽게 설명해.
• '사용자'라는 표현 대신 '유저'를 사용해.

출력 예시

All right everybody it is now time to set up our environment
좋아 모두, 이제 environment를 설정할 시간이야.

and not to set up one environment
근데 딱 하나만 설정하는 게 아니야.

I just want to show you how to set up a environment
내가 보여주려는 건 environment 만드는 방법이야.

because the way we're going to work in this course is by dividing and conquering okay
이번 코스는 divide and conquer 방식으로 진행할 거거든:"""


def split_by_sentences_with_timing(transcription_segments, max_duration=30):
    """
    Split transcription into subtitle-appropriate chunks based on timing and sentences
    max_duration: maximum seconds per subtitle chunk
    """
    chunks = []
    current_chunk = {"text": "", "start": None, "end": None}

    for segment in transcription_segments:
        segment_duration = segment["end"] - segment["start"]

        # If adding this segment would make chunk too long, finish current chunk
        if (
            current_chunk["start"] is not None
            and segment["end"] - current_chunk["start"] > max_duration
        ):

            if current_chunk["text"].strip():
                chunks.append(current_chunk)

            current_chunk = {
                "text": segment["text"].strip(),
                "start": segment["start"],
                "end": segment["end"],
            }
        else:
            # Add to current chunk
            if current_chunk["start"] is None:
                current_chunk["start"] = segment["start"]
            current_chunk["end"] = segment["end"]
            current_chunk["text"] += " " + segment["text"].strip()

    # Add final chunk
    if current_chunk["text"].strip():
        chunks.append(current_chunk)

    return chunks


def translate_chunk_with_gpt(text_chunk, chunk_index):
    """
    Translate a text chunk using GPT-4o-mini
    """
    try:
        print(f"[INFO] Translating chunk {chunk_index + 1}...")

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": TRANSLATION_PROMPT},
                {"role": "user", "content": text_chunk},
            ],
            temperature=0.3,
            max_tokens=15000,
        )

        translation = response.choices[0].message.content
        print(f"[INFO] Chunk {chunk_index + 1} translation complete.")
        return translation

    except Exception as e:
        print(f"[ERROR] Translation failed for chunk {chunk_index + 1}: {e}")
        # Fallback: return original text with basic formatting
        lines = text_chunk.split(".")
        formatted = ""
        for line in lines:
            if line.strip():
                formatted += f"**{line.strip()}**\n번역 실패\n\n"
        return formatted


def transcribe_and_translate_audio(file_path, output_directory, subtitle_duration=30):
    """
    Main function that transcribes and translates in one process
    subtitle_duration: target duration for each subtitle chunk in seconds
    """
    # Transcribe directly with word-level timestamps
    print(f"[INFO] Transcribing '{file_path}' directly with timestamps...")
    result = whisper_model.transcribe(file_path, word_timestamps=True)

    # Split into subtitle-appropriate chunks
    print(f"[INFO] Splitting transcription into {subtitle_duration}-second chunks...")
    chunks = split_by_sentences_with_timing(result["segments"], subtitle_duration)
    print(f"[INFO] Created {len(chunks)} subtitle chunks.")

    # Translate each chunk
    full_translated_content = ""
    for index, chunk in enumerate(chunks):
        print(
            f"[INFO] Processing chunk {index + 1}/{len(chunks)} ({chunk['start']:.1f}s - {chunk['end']:.1f}s)"
        )

        # Clean up the text
        clean_text = chunk["text"].strip()
        if not clean_text:
            continue

        # Translate
        translated_chunk = translate_chunk_with_gpt(clean_text, index)

        # Add timing info and content
        full_translated_content += (
            f"\n[Chunk {index + 1}] ({chunk['start']:.1f}s - {chunk['end']:.1f}s)\n"
        )
        full_translated_content += translated_chunk + "\n\n"

        # Small delay to avoid rate limiting
        time.sleep(0.5)

    return full_translated_content


def save_to_docx(content, output_file):
    print(f"[INFO] Saving to '{output_file}'...")
    doc = Document()

    bold_re = re.compile(r"^\s*\*\*(.+?)\*\*\s*$")  # 앞뒤 공백/개행 허용, **...** 캡처

    # 줄 단위 순회 (개행·공백 정리)
    lines = content.splitlines()

    for raw in lines:
        line = raw.rstrip()  # 오른쪽 공백 제거 (Markdown의 '두 칸 공백+개행' 무력화)
        if not line:
            continue

        # [Chunk N] 타이밍 헤더는 그대로 추가
        if line.startswith("[Chunk "):
            doc.add_paragraph(line)
            continue

        # **영어** 패턴이면 굵게
        m = bold_re.match(line)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(1))
            run.bold = True
            continue

        # 혹시 앞은 **로 시작하지만 끝에 **가 빠진 사례까지 케어
        if line.lstrip().startswith("**"):
            # 맨 앞 ** 제거 후, 닫는 **가 있으면 잘라냄
            stripped = line.lstrip()[2:]
            end = stripped.find("**")
            text = stripped[:end] if end != -1 else stripped
            p = doc.add_paragraph()
            run = p.add_run(text.strip())
            run.bold = True
            continue

        # 일반 번역 줄
        doc.add_paragraph(line)

    doc.save(output_file)
    print(f"[INFO] File saved successfully!")


# Main execution
if __name__ == "__main__":
    directory = r"C:\Users\Jonathan\Documents\kbdocs\text_result"
    media_files = [f for f in os.listdir(directory) if f.endswith((".mp4", ".mp3"))]

    if not media_files:
        print("[WARNING] No media files found in the directory.")
    else:
        for media_file in media_files:
            audio_file_path = os.path.join(directory, media_file)
            print(f"[INFO] Processing file: {audio_file_path}")

            # Process with 25-second subtitle chunks (adjust as needed)
            translated_content = transcribe_and_translate_audio(
                audio_file_path, directory, subtitle_duration=25
            )

            print("[INFO] Transcription and translation completed!")

            # Save to file
            output_file_name = os.path.splitext(media_file)[0] + "_translated.docx"
            output_file_path = os.path.join(directory, output_file_name)
            save_to_docx(translated_content, output_file_path)

            print(f"[INFO] Final output saved to: {output_file_path}")
