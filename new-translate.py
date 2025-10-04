import os
from docx import Document
import whisper
from dotenv import load_dotenv
import openai
from openai import OpenAI
import re
import time

# ============================================================================
# 환경 설정 및 초기화
# ============================================================================

# .env 파일에서 환경 변수 로드 (OPENAI_API_KEY 등)
load_dotenv()

# OpenAI 클라이언트 초기화 (GPT API 사용을 위해)
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Whisper 모델 로드 (음성을 텍스트로 변환하는 AI 모델)
print("[INFO] Loading Whisper model (medium.en)...")
whisper_model = whisper.load_model("medium.en", device="cpu")
print("[INFO] Whisper model loaded successfully!")


# ============================================================================
# GPT 번역 프롬프트 (시스템 메시지)
# ============================================================================

# GPT에게 어떻게 번역할지 지시하는 상세한 한국어 프롬프트
# 니콜라스 강의 스타일에 맞춰 캐주얼하게 번역하고, 전문 용어는 영어로 유지
TRANSLATION_PROMPT = """너는 노마드코더(Nomad Coders)의 니콜라스 강의를 번역하는 역할이야.

니콜라스의 스타일은 다음과 같아:
• 캐주얼하고 재미있는 말투
• 실용적이고 직관적인 설명
• 격식 없는 자연스러운 대화체

번역 규칙:
• 입력된 transcript는 **짧은 구간**으로 주어져. 이걸 **1~2줄의 자막 단위로만** 끊어서 정리해.
• **절대 3문장 이상을 하나로 묶지 마.** 한 영어 구절당 1~2문장 최대.
• 영어 문장은 절대 수정하지 않고 그대로 두고, 굵게 표시 (bold).
• 바로 아래에 한국어 번역을 붙여. 한국어 번역은 자연스럽고 캐주얼하게, "해, 할 거야, 하는 거지, 그렇고" 같은 대화체로 작성해.
• 헤딩, 구분선 같은 건 넣지 말고 깔끔하게 이어서 작성해.
• you라는 단어 나오면 "여러분", "여러분이", 아니면 문맥상 자연스러우면 여러분 넣지 않고 번역해.
• 전문 용어(import, function, class, async/await, coroutine, decorator, variable, API, endpoint, JSON, YAML, pip, venv, package manager, CrewAI, AutoGen, OpenAI Agents SDK, LangGraph, Google Agent Builder, kit, LLM, prompt, system message, user message, state, tool, orchestrator 등)는 번역하지 않고 그대로 둬.
• snake_case / camelCase / kebab-case / SNAKE_CASE 등은 그대로 유지 (예: transfer_agent, web_search_tool, job_search_agent).
• agent라고만 딱 나오면 한국말로 "에이전트"로 쓰고, 다른 전문용어는 그대로 써.
• 코드 블록(python, bash, json 등)은 번역하지 않고 그대로 둬.
• 특정한 행동이나 작성 지시가 포함된 경우 (예: "위 같이 작성", "화면 참조", "추가사항 적용")는 문맥에 맞게 간결하고 자연스럽게 설명해.
• "사용자"라는 표현 대신 항상 "유저"라고 번역해.

중요 (VERY IMPORTANT):
• 너무 흔한 단어들(state, function calling, import, class 등)은 그대로 두고 (한국말) 붙이지 않아.
• 강의에서 처음 나오거나 헷갈릴 수 있는 개념은 → 처음 등장할 때만 짧고 직관적인 한국어 설명을 () 안에 넣어줘.
  예:
  - Conditional_Edge(조건에 따라 다음 Node로 이동 여부를 결정하는 엣지)
  - Edge(엣지, 노드 간 연결선)
  - Node(노드, 작업이나 상태를 담는 단위 박스)
  - Decorator(데코레이터, 함수/클래스에 기능을 쉽게 덧붙이는 문법 도구)
  이후 다시 나올 땐 영어만 써.

**줄바꿈 강제 규칙 (CRITICAL):**
• 접속사 기준: and, but, so, because, or 등이 나오면 그 앞에서 끊어.
• 구두점 기준: 마침표(.), 쉼표(,), 세미콜론(;) 뒤에서 끊어.
• 의미 단위: 한 영어 구절이 80자 넘어가면 무조건 더 쪼개.
• 출력 형식: 반드시 "영어 1~2문장 → 한국어 1~2문장" 페어. 3문장 이상 절대 금지.

출력 예시:
**All right everybody it is now time to set up our environment**
좋아 모두, 이제 environment(환경)를 설정할 시간이야.

**and not to set up one environment**
근데 딱 하나만 설정하는 건 아니야.

**I just want to show you how to set up a environment**
내가 보여주려는 건 environment 만드는 방법이야.

**because the way we're going to work in this course is by dividing and conquering okay**
왜냐면 이번 코스는 divide and conquer(나눠서 정복하기) 방식으로 진행할 거거든.
"""


# ============================================================================
# 함수 1: 전체 텍스트를 작은 청크(덩어리)로 나누기
# ============================================================================


def smart_split_transcript(text, max_chars=200):
    """
    전체 transcript를 자연스러운 문장 단위로 나누는 함수

    목적:
    - GPT에게 한 번에 너무 긴 텍스트를 주면 3문장 이상으로 묶어서 번역함
    - 이를 방지하기 위해 미리 작은 단위(청크)로 쪼갬

    동작 방식:
    1. 마침표(.), 물음표(?), 느낌표(!) 기준으로 문장 분리
    2. 각 문장을 합쳐가며 max_chars(기본 200자) 이내로 청크 생성
    3. 200자 넘으면 새로운 청크 시작

    파라미터:
    - text: Whisper가 추출한 전체 영어 텍스트
    - max_chars: 청크 하나당 최대 글자 수 (작을수록 더 짧게 분할)

    반환값:
    - chunks: 문자열 리스트 (예: ["첫 번째 청크...", "두 번째 청크..."])
    """
    # 정규식으로 문장 분리: 마침표+공백, 물음표+공백, 느낌표+공백 기준
    # ([.!?]+\s+) 패턴은 구두점과 공백을 함께 캡처해서 나중에 다시 붙일 수 있게 함
    sentences = re.split(r"([.!?]+\s+)", text)

    chunks = []  # 최종 청크들을 저장할 리스트
    current_chunk = ""  # 현재 만들고 있는 청크

    # 문장을 2개씩 묶어서 처리 (문장 + 구두점)
    for i in range(0, len(sentences), 2):
        sentence = sentences[i]  # 실제 문장 내용
        delimiter = sentences[i + 1] if i + 1 < len(sentences) else ""  # 구두점 부분

        full_sentence = sentence + delimiter  # 문장 + 구두점 합치기

        # 현재 청크에 이 문장을 추가하면 max_chars를 초과하는지 체크
        if current_chunk and len(current_chunk) + len(full_sentence) > max_chars:
            # 초과하면: 현재 청크를 저장하고, 새 청크 시작
            chunks.append(current_chunk.strip())
            current_chunk = full_sentence
        else:
            # 초과 안 하면: 현재 청크에 계속 추가
            current_chunk += full_sentence

    # 마지막 남은 청크 저장
    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks


# ============================================================================
# 함수 2: 청크 하나를 GPT로 번역하기 (이전 문맥 포함)
# ============================================================================


def translate_chunk_with_context(chunk, previous_context=""):
    """
    작은 청크 하나를 GPT-4o-mini로 번역하는 함수

    목적:
    - 청크별로 번역하면서도 이전 내용과의 연결성 유지
    - 이전 청크의 마지막 부분을 "참고용"으로 GPT에게 제공

    동작 방식:
    1. 이전 청크의 마지막 100자를 "문맥"으로 함께 전달
    2. GPT에게 현재 청크만 번역하도록 지시
    3. 번역 결과 반환

    파라미터:
    - chunk: 번역할 현재 청크 (영어 텍스트)
    - previous_context: 이전 청크의 내용 (문맥 유지용)

    반환값:
    - translation: GPT가 번역한 결과 (영어 bold + 한국어 형식)
    """
    try:
        # 이전 문맥이 있으면 함께 전달 (마지막 100자만 사용)
        if previous_context:
            user_message = f"[Previous context for continuity: {previous_context[-100:]}]\n\n[Translate this]: {chunk}"
        else:
            user_message = chunk

        # GPT API 호출
        response = client.chat.completions.create(
            model="gpt-4o-mini",  # 사용할 모델
            messages=[
                {"role": "system", "content": TRANSLATION_PROMPT},  # 번역 지시사항
                {"role": "user", "content": user_message},  # 번역할 내용
            ],
            temperature=0.2,  # 낮을수록 일관된 번역 (0~2 범위)
            max_tokens=1500,  # 청크가 작으니 1500 토큰이면 충분
        )

        # GPT 응답에서 번역 결과 추출
        translation = response.choices[0].message.content
        return translation

    except Exception as e:
        # API 오류 발생 시 대체 처리
        print(f"[ERROR] Translation failed for chunk: {e}")

        # 폴백: 마침표로 나눠서 기본 포맷으로 반환
        lines = chunk.split(".")
        formatted = ""
        for line in lines:
            if line.strip():
                formatted += f"**{line.strip()}**\n번역 실패\n\n"
        return formatted


# ============================================================================
# 함수 3: 오디오 파일을 transcribe하고 청크별로 번역하기 (메인 로직)
# ============================================================================


def transcribe_and_translate_audio(file_path, output_directory, chunk_max_chars=200):
    """
    전체 파이프라인을 실행하는 메인 함수

    처리 순서:
    1. Whisper로 오디오 → 영어 텍스트 변환 (transcription)
    2. 전체 텍스트를 작은 청크들로 분할
    3. 각 청크를 GPT로 번역 (이전 청크 문맥 포함)
    4. 모든 번역 결과를 하나로 합쳐서 반환

    파라미터:
    - file_path: 처리할 오디오/비디오 파일 경로
    - output_directory: 출력 파일을 저장할 디렉토리 (현재는 미사용)
    - chunk_max_chars: 청크 하나당 최대 글자 수 (기본 200)
                       작게 할수록 더 짧은 단위로 번역됨 (150~300 추천)

    반환값:
    - final_content: 모든 청크의 번역 결과를 합친 최종 문자열
    """
    # 1단계: Whisper로 음성 인식 (transcription)
    print(f"[INFO] Transcribing '{file_path}'...")
    result = whisper_model.transcribe(file_path, word_timestamps=True)

    # 전체 텍스트 추출
    full_text = result["text"].strip()

    # 텍스트가 없으면 종료
    if not full_text:
        print("[WARNING] No transcription found.")
        return ""

    print(f"[INFO] Full transcription length: {len(full_text)} characters")

    # 2단계: 전체 텍스트를 청크로 분할
    chunks = smart_split_transcript(full_text, max_chars=chunk_max_chars)
    print(f"[INFO] Split into {len(chunks)} chunks for translation")

    # 3단계: 각 청크를 순서대로 번역
    all_translations = []  # 모든 번역 결과를 저장할 리스트
    previous_chunk = ""  # 이전 청크 내용 (문맥 유지용)

    for i, chunk in enumerate(chunks):
        print(f"[INFO] Translating chunk {i+1}/{len(chunks)}...")

        # 현재 청크 번역 (이전 청크 문맥 포함)
        translation = translate_chunk_with_context(
            chunk, previous_context=previous_chunk
        )
        all_translations.append(translation)

        # 다음 번역을 위해 현재 청크 저장
        previous_chunk = chunk

        # API rate limit 방지를 위한 짧은 대기
        time.sleep(0.5)

    # 4단계: 모든 번역 결과를 합치기
    # \n\n으로 청크 사이를 구분 (빈 줄 2개)
    final_content = "\n\n".join(all_translations)

    print("[INFO] All chunks translated and combined.")
    return final_content


# ============================================================================
# 함수 4: 번역 결과를 Word 문서(.docx)로 저장하기
# ============================================================================


def save_to_docx_improved(content, output_file):
    """
    번역된 텍스트를 Word 문서로 저장하는 함수

    포맷 규칙:
    - **텍스트** 형식 → Bold 처리 (영어 원문)
    - 일반 텍스트 → 평문 처리 (한국어 번역)
    - 빈 줄 → 문단 간격

    파라미터:
    - content: 번역된 전체 텍스트 (GPT 출력 결과)
    - output_file: 저장할 파일 경로 (예: "lecture_translated.docx")

    반환값: 없음 (파일 저장만 수행)
    """
    print(f"[INFO] Saving to '{output_file}'...")

    # 새 Word 문서 생성
    doc = Document()

    # 줄 단위로 분리
    lines = content.split("\n")

    for line in lines:
        line = line.strip()  # 앞뒤 공백 제거

        # 빈 줄이면 문단 간격 추가
        if not line:
            doc.add_paragraph()
            continue

        # **로 감싸진 줄 = 영어 원문 (Bold 처리)
        bold_match = re.match(r"\*\*(.*?)\*\*", line)
        if bold_match:
            english_text = bold_match.group(1).strip()
            if english_text:
                p = doc.add_paragraph()
                run = p.add_run(english_text)
                run.bold = True  # 굵게
        else:
            # 일반 텍스트 = 한국어 번역 (평문)
            doc.add_paragraph(line)

    # 파일 저장
    doc.save(output_file)
    print(f"[INFO] File saved successfully!")


# ============================================================================
# 메인 실행부: 디렉토리의 모든 미디어 파일 처리
# ============================================================================

if __name__ == "__main__":
    # 처리할 파일들이 있는 디렉토리 경로
    directory = r"C:\Users\Jonathan\Documents\kbdocs\text_result"

    # 디렉토리에서 .mp4, .mp3 파일만 찾기
    media_files = [f for f in os.listdir(directory) if f.endswith((".mp4", ".mp3"))]

    # 파일이 없으면 경고 출력
    if not media_files:
        print("[WARNING] No media files found in the directory.")
    else:
        # 각 파일을 순서대로 처리
        for media_file in media_files:
            audio_file_path = os.path.join(directory, media_file)
            print(f"[INFO] Processing file: {audio_file_path}")

            # 핵심: 청크 기반 번역 실행
            # chunk_max_chars 조정 가능:
            # - 150: 매우 짧은 단위로 분할 (번역 더 짧아짐)
            # - 200: 적당한 길이 (기본값)
            # - 300: 조금 더 긴 문맥 유지
            translated_content = transcribe_and_translate_audio(
                audio_file_path,
                directory,
                chunk_max_chars=200,  # 이 값을 조절해서 세밀도 조정
            )

            # 번역 결과가 있으면 파일 저장
            if translated_content:
                print("[INFO] Transcription and translation completed!")

                # 출력 파일명 생성 (원본 파일명 + _translated.docx)
                output_file_name = os.path.splitext(media_file)[0] + "_translated.docx"
                output_file_path = os.path.join(directory, output_file_name)

                # Word 문서로 저장
                save_to_docx_improved(translated_content, output_file_path)

                print(f"[INFO] Final output saved to: {output_file_path}")
            else:
                print("[WARNING] No content to save.")
