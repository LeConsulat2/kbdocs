import os
from pydub import AudioSegment
import whisper
from docx import Document
from dotenv import load_dotenv
import openai
import re
import nltk
import pysrt # SRT 파일 생성을 위한 라이브러리
from datetime import timedelta # 시간 계산을 위한 라이브러리

# NLTK data 다운로드 (한 번만 실행)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# .env 파일 로드
load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')

# Whisper 모델 로드
# 성능과 속도를 고려하여 'medium.en' 모델을 사용합니다.
# 더 빠른 처리가 필요하면 'small.en' 또는 'base.en'을 고려할 수 있습니다.
print("[INFO] Loading Whisper model (medium.en)... This may take a while.")
try:
    whisper_model = whisper.load_model("medium.en", device="cpu") # GPU 사용 가능하면 'cuda'로 변경
    print("[INFO] Whisper model loaded successfully!")
except Exception as e:
    print(f"[ERROR] Failed to load Whisper model: {e}")
    print("[ERROR] Please ensure you have downloaded the Whisper models correctly.")
    exit()

def format_timestamp(milliseconds):
    """밀리초를 'HH:MM:SS,ms' 형식으로 변환"""
    td = timedelta(milliseconds=milliseconds)
    hours, remainder = divmod(td.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    ms = td.microseconds // 1000
    return f"{hours:02}:{minutes:02}:{seconds:02},{ms:03}"

def create_srt_from_segments(segments, output_filepath):
    """Whisper 세그먼트 정보를 바탕으로 SRT 파일 생성"""
    subs = pysrt.SubRipFile()
    for i, segment in enumerate(segments):
        start_ms = int(segment['start'] * 1000)
        end_ms = int(segment['end'] * 1000)
        text = segment['text'].strip()

        # 빈 텍스트는 건너뜁니다.
        if not text:
            continue

        sub = pysrt.SubRipItem(
            index=i + 1,
            start=pysrt.SubRipTime(milliseconds=start_ms),
            end=pysrt.SubRipTime(milliseconds=end_ms),
            text=text
        )
        subs.append(sub)
    
    try:
        subs.save(output_filepath, encoding='utf-8')
        print(f"[INFO] SRT file saved to '{output_filepath}'")
    except Exception as e:
        print(f"[ERROR] Failed to save SRT file: {e}")

def create_bilingual_format(english_text_segments):
    """
    영어 세그먼트 (텍스트와 타임스탬프 정보 포함) 리스트를 받아
    이중 언어 형식으로 변환 (영어 Bold + 한국어 번역)
    """
    if not openai.api_key:
        print("[WARNING] No OpenAI API key found. Creating English-only format.")
        # API 키가 없으면 영어 텍스트만 볼드로 반환
        formatted = []
        for segment in english_text_segments:
            if segment['text'].strip():
                formatted.append(f"**{segment['text'].strip()}**")
        return '\n\n'.join(formatted)
    
    translation_prompt = """You are an excellent translator who translates Nicolas's Nomad Coders lectures.

Translation Rules:
• DO NOT modify the input English sentences. Keep them as they are.
• The English sentence should be **bolded**, followed immediately by its Korean translation on the next line.
• Translate in a casual, natural, conversational tone (e.g., 해, 할 거야, 하는 거지, 그렇고).
• Keep technical terms as they are (e.g., console, deploy, import, variable, function).
• Convert 'Profile ID' to 'profile_id' format.
• Use '유저' instead of '사용자'.
• Maintain appropriate spacing between each paragraph.

Output Format:
**Original English text**
Korean translation

Translate the following text according to the rules above:"""

    bilingual_parts = []
    
    # 세그먼트들을 문단 단위로 묶기 위한 임시 저장소
    current_paragraph_english_segments = []
    
    # 세그먼트들을 적절한 단위로 묶어 번역 요청
    # 여기서는 약 2~3문장 정도 또는 최대 150-200자 정도를 기준으로 묶습니다.
    # Whisper의 segment는 이미 의미 단위로 잘 나뉘어 있으므로, 이를 활용합니다.
    
    for i, segment in enumerate(english_text_segments):
        current_paragraph_english_segments.append(segment)
        current_text = " ".join([s['text'] for s in current_paragraph_english_segments])

        # 문장 끝 (마침표, 물음표 등) 또는 적당한 길이일 때 번역 요청
        # 너무 짧은 세그먼트는 다음 세그먼트와 합쳐서 처리
        if len(current_text) > 50 and (current_text.strip().endswith(('.', '?', '!')) or i == len(english_text_segments) - 1):
            
            text_to_translate = " ".join([s['text'].strip() for s in current_paragraph_english_segments if s['text'].strip()])
            
            if not text_to_translate.strip():
                current_paragraph_english_segments = []
                continue

            try:
                print(f"[INFO] Translating chunk (segment {i - len(current_paragraph_english_segments) + 2} to {i + 1})...")
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": translation_prompt},
                        {"role": "user", "content": text_to_translate}
                    ],
                    temperature=0.2
                )
                bilingual_parts.append(response.choices[0].message.content.strip())
            except Exception as e:
                print(f"[ERROR] Translation failed for chunk: {e}")
                # Fallback: create manual bilingual format
                fallback_format = []
                for seg in current_paragraph_english_segments:
                    if seg['text'].strip():
                        fallback_format.append(f"**{seg['text'].strip()}**")
                        fallback_format.append("(번역 실패)")
                bilingual_parts.append('\n\n'.join(fallback_format))
            
            current_paragraph_english_segments = [] # 번역 후 초기화
    
    # 마지막 남은 세그먼트 처리 (루프가 끝났는데 번역되지 않은 텍스트가 있을 경우)
    if current_paragraph_english_segments:
        text_to_translate = " ".join([s['text'].strip() for s in current_paragraph_english_segments if s['text'].strip()])
        if text_to_translate.strip():
            try:
                print(f"[INFO] Translating final chunk...")
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "system", "content": translation_prompt},
                        {"role": "user", "content": text_to_translate}
                    ],
                    temperature=0.2
                )
                bilingual_parts.append(response.choices[0].message.content.strip())
            except Exception as e:
                print(f"[ERROR] Translation failed for final chunk: {e}")
                fallback_format = []
                for seg in current_paragraph_english_segments:
                    if seg['text'].strip():
                        fallback_format.append(f"**{seg['text'].strip()}**")
                        fallback_format.append("(번역 실패)")
                bilingual_parts.append('\n\n'.join(fallback_format))

    return '\n\n'.join(bilingual_parts)

# 1. 오디오 파일을 WAV로 변환
def convert_to_wav(file_path, output_directory):
    """지원되는 모든 오디오/비디오 파일을 WAV로 변환"""
    print(f"[INFO] Converting '{file_path}' to WAV format...")
    try:
        audio = AudioSegment.from_file(file_path)
        wav_file_path = os.path.join(
            output_directory, os.path.splitext(os.path.basename(file_path))[0] + ".wav"
        )
        audio.export(wav_file_path, format="wav")
        print(f"[INFO] Conversion complete. Saved as '{wav_file_path}'")
        return wav_file_path
    except Exception as e:
        print(f"[ERROR] Failed to convert audio: {e}")
        return None

# 2. Whisper로 전사 (세그먼트 정보 활용)
def transcribe_audio_with_segments(audio_file_path):
    """
    오디오 파일을 전사하고, 세그먼트별 텍스트와 타임스탬프를 반환
    """
    print(f"[INFO] Transcribing audio from '{audio_file_path}'...")
    try:
        result = whisper_model.transcribe(
            audio_file_path,
            word_timestamps=False, # 단어 단위는 SRT에서 정확히 필요하지 않으므로 False로 설정 (성능 향상)
            verbose=False # 진행 상황 메시지 억제
        )
        # result['segments']는 각 세그먼트의 'start', 'end', 'text'를 포함
        print(f"[INFO] Transcription complete. Found {len(result.get('segments', []))} segments.")
        return result.get('segments', [])
    except Exception as e:
        print(f"[ERROR] Transcription failed: {e}")
        return []

# 3. 향상된 DOCX 저장
def save_text_to_docx(text, output_file):
    """텍스트를 DOCX 파일로 저장하며, **볼드** 형식 처리"""
    print(f"[INFO] Saving to '{output_file}'...")
    doc = Document()
    
    # 스타일 설정 (필요시 추가)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = doc.styles['Normal'].font.size
    
    lines = text.split('\n')
    for line in lines:
        paragraph = doc.add_paragraph()
        if line.strip(): # 빈 줄이 아니면 처리
            # **볼드** 텍스트 처리
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)
        else:
            doc.add_paragraph() # 빈 줄을 DOCX에서도 빈 단락으로 유지
    
    try:
        doc.save(output_file)
        print(f"[INFO] Saved successfully to '{output_file}'!")
    except Exception as e:
        print(f"[ERROR] Failed to save DOCX file: {e}")


# 4. 메인 처리 함수
def process_media_file(file_path, output_directory, create_bilingual=True, create_srt=True):
    """
    단일 미디어 파일을 처리하여 전사, 번역, DOCX, SRT 파일 생성
    """
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    
    # 1. 오디오 파일을 WAV로 변환 (Whisper 입력용)
    wav_file_path = convert_to_wav(file_path, output_directory)
    if not wav_file_path:
        return

    # 2. Whisper로 전사 및 세그먼트 정보 획득
    segments = transcribe_audio_with_segments(wav_file_path)
    
    # 임시 WAV 파일 삭제
    os.remove(wav_file_path)
    print(f"[INFO] Cleaned up temporary WAV file: {wav_file_path}")

    if not segments:
        print(f"[WARNING] No transcription segments found for '{file_path}'. Skipping further processing.")
        return

    # 3. SRT 파일 생성 (타임스탬프 기반의 자막)
    if create_srt:
        srt_output_path = os.path.join(output_directory, f"{base_name}.srt")
        create_srt_from_segments(segments, srt_output_path)
    
    # 4. 이중 언어 형식 생성 및 DOCX 저장
    if create_bilingual:
        print("[INFO] Creating bilingual format...")
        # create_bilingual_format 함수가 세그먼트 리스트를 받도록 수정
        bilingual_text = create_bilingual_format(segments) 
        
        if bilingual_text:
            bilingual_docx_path = os.path.join(output_directory, f"{base_name}_bilingual.docx")
            save_text_to_docx(bilingual_text, bilingual_docx_path)
            print(f"[INFO] Bilingual DOCX file saved: {bilingual_docx_path}")
        else:
            print("[WARNING] No bilingual content generated.")
    else:
        # 이중 언어 생성을 원치 않으면, 원문 텍스트를 DOCX로 저장
        full_transcript_text = "\n\n".join([s['text'].strip() for s in segments if s['text'].strip()])
        if full_transcript_text:
            transcript_docx_path = os.path.join(output_directory, f"{base_name}_transcript.docx")
            save_text_to_docx(full_transcript_text, transcript_docx_path)
            print(f"[INFO] English transcript DOCX file saved: {transcript_docx_path}")
        else:
            print("[WARNING] No English transcript content generated.")


# 메인 실행부
if __name__ == "__main__":
    print("\n" + "=" * 50)
    print("      Nomad Coders Lecture Transcriber & Translator      ")
    print("=" * 50 + "\n")

    # 사용자로부터 파일 또는 디렉토리 경로 입력받기
    input_path = input("Enter the path to the media file or directory containing media files: ").strip()

    if not input_path:
        print("[ERROR] No path entered. Exiting.")
        exit()

    # 출력 디렉토리는 입력 디렉토리와 동일하게 설정
    output_dir = os.path.dirname(input_path) if os.path.isfile(input_path) else input_path
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"[INFO] Created output directory: {output_dir}")

    CREATE_BILINGUAL = True  # 이중 언어 형식 생성 (영어 볼드 + 한국어)
    CREATE_SRT = True        # SRT 자막 파일 생성

    media_files_to_process = []
    if os.path.isfile(input_path):
        media_files_to_process.append(input_path)
    elif os.path.isdir(input_path):
        media_files_to_process = [
            os.path.join(input_path, f)
            for f in os.listdir(input_path)
            if f.lower().endswith((".mp4", ".mp3", ".wav"))
        ]
    else:
        print(f"[ERROR] Invalid path: '{input_path}'. Please provide a valid file or directory.")
        exit()

    if not media_files_to_process:
        print("[WARNING] No supported media files found in the specified path.")
    else:
        for i, media_file_path in enumerate(media_files_to_process):
            print(f"\n[INFO] Processing file {i+1}/{len(media_files_to_process)}: {os.path.basename(media_file_path)}")
            print("=" * 60)
            process_media_file(media_file_path, output_dir, CREATE_BILINGUAL, CREATE_SRT)
            print(f"[INFO] Finished processing: {os.path.basename(media_file_path)}")
            print("=" * 60)

    print("\nAll tasks completed!")