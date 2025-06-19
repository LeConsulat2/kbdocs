import os
from pydub import AudioSegment
import whisper
from docx import Document
from dotenv import load_dotenv
import textwrap
import openai
import re
import nltk
from nltk.tokenize import sent_tokenize
import spacy

# Download required NLTK data (run once)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# Load spaCy model for better sentence detection
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("[WARNING] spaCy English model not found. Install with: python -m spacy download en_core_web_sm")
    nlp = None

# .env 파일 로드
load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')

# Whisper 모델 로드
print("[INFO] Loading Whisper model (medium.en)...")
whisper_model = whisper.load_model("medium.en", device="cpu")
print("[INFO] Whisper model loaded successfully!")


def detect_speech_patterns(text):
    """
    Detect natural speech patterns and breaks using multiple methods
    """
    # Common speech fillers and transition words
    speech_markers = [
        "okay", "alright", "all right", "so", "now", "well", "um", "uh",
        "you know", "like", "basically", "actually", "really", "right",
        "perfect", "fantastic", "good", "great", "yes", "yep", "yeah",
        "let's", "we're going to", "we are going to", "i'm going to",
        "but", "and", "then", "next", "first", "second", "also",
        "however", "because", "since", "when", "where", "what", "how"
    ]
    
    # Technical action words that often indicate new steps
    action_words = [
        "click", "press", "select", "choose", "copy", "paste", "run",
        "create", "add", "remove", "delete", "save", "open", "close",
        "go to", "navigate", "scroll", "drag", "drop", "install"
    ]
    
    return speech_markers, action_words


def smart_sentence_break(text, max_line_length=80):
    """
    Intelligently break text into natural speech segments
    """
    speech_markers, action_words = detect_speech_patterns(text)
    
    # First, split by common sentence endings
    sentences = re.split(r'[.!?]+', text)
    formatted_lines = []
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
            
        # Check if sentence is too long and needs breaking
        if len(sentence) > max_line_length:
            # Try to break at natural speech points
            words = sentence.split()
            current_line = []
            
            for i, word in enumerate(words):
                current_line.append(word)
                current_text = ' '.join(current_line)
                
                # Check for natural break points
                word_lower = word.lower().rstrip('.,!?')
                
                # Break after speech markers or action words
                if (word_lower in speech_markers or 
                    any(action in current_text.lower() for action in action_words)):
                    
                    # Don't break if line is too short
                    if len(current_text) > 30:
                        formatted_lines.append(current_text)
                        current_line = []
                        continue
                
                # Break if line is getting too long
                if len(current_text) > max_line_length:
                    # Try to break at a natural point
                    if len(current_line) > 1:
                        # Keep the last word for next line
                        line_to_add = ' '.join(current_line[:-1])
                        formatted_lines.append(line_to_add)
                        current_line = [current_line[-1]]
                    else:
                        # Force break if word is too long
                        formatted_lines.append(current_text)
                        current_line = []
            
            # Add remaining words
            if current_line:
                formatted_lines.append(' '.join(current_line))
        else:
            formatted_lines.append(sentence)
    
    return formatted_lines


def advanced_format_transcription(transcription):
    """
    Advanced formatting that mimics natural speech patterns
    """
    # Remove extra whitespace and normalize
    text = re.sub(r'\s+', ' ', transcription.strip())
    
    # Get smart sentence breaks
    lines = smart_sentence_break(text)
    
    # Post-process for better formatting
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Capitalize first letter if needed
        if line and line[0].islower():
            line = line[0].upper() + line[1:]
        
        # Add proper punctuation if missing
        if line and line[-1] not in '.!?':
            line += '.'
        
        formatted_lines.append(line)
    
    # Join with double line breaks for readability
    return '\n\n'.join(formatted_lines)


def create_bilingual_format(english_text):
    """
    Create bilingual format with English (bold) followed by Korean translation
    """
    if not openai.api_key:
        print("[WARNING] No OpenAI API key found. Creating English-only format.")
        # Return English text formatted with bold markers
        paragraphs = english_text.split('\n\n')
        formatted = []
        for paragraph in paragraphs:
            if paragraph.strip():
                formatted.append(f"**{paragraph.strip()}**")
        return '\n\n'.join(formatted)
    
    translation_prompt = """너는 노마드코더(Nomad Coders)의 니콜라스 강의를 번역하는 역할이야. 

번역 규칙:
• 입력된 영어 문장은 절대 수정하지 않고 그대로 둬
• 영어 문장은 **굵게** 표시하고, 바로 아래에 한국어 번역
• 캐주얼하고 자연스러운 대화체로 번역 (해, 할 거야, 하는 거지, 그렇고)
• 전문 용어는 원문 그대로 (console, deploy, import, variable, function 등)
• Profile ID → profile_id 형식으로 변환
• '사용자' 대신 '유저' 사용
• 각 문단 사이에 적절한 공백 유지

출력 형식:
**영어 원문**
한국어 번역

다음 텍스트를 위 형식으로 번역해:"""

    # Process in chunks
    paragraphs = english_text.split('\n\n')
    bilingual_parts = []
    
    chunk_size = 5  # 5개 문단씩 처리
    for i in range(0, len(paragraphs), chunk_size):
        chunk_paragraphs = paragraphs[i:i+chunk_size]
        chunk_text = '\n\n'.join([p for p in chunk_paragraphs if p.strip()])
        
        if not chunk_text.strip():
            continue
            
        try:
            print(f"[INFO] Translating chunk {i//chunk_size + 1}/{(len(paragraphs)-1)//chunk_size + 1}...")
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": translation_prompt},
                    {"role": "user", "content": chunk_text}
                ],
                temperature=0.2
            )
            bilingual_parts.append(response.choices[0].message.content)
        except Exception as e:
            print(f"[ERROR] Translation failed for chunk: {e}")
            # Fallback: create manual bilingual format
            fallback_format = []
            for paragraph in chunk_paragraphs:
                if paragraph.strip():
                    fallback_format.append(f"**{paragraph.strip()}**")
                    fallback_format.append("(번역 실패)")
            bilingual_parts.append('\n\n'.join(fallback_format))
    
    return '\n\n'.join(bilingual_parts)


# 1. 오디오 파일을 WAV로 변환
def convert_to_wav(file_path, output_directory):
    print(f"[INFO] Converting '{file_path}' to WAV format...")
    audio = AudioSegment.from_file(file_path)
    wav_file_path = os.path.join(
        output_directory, os.path.splitext(os.path.basename(file_path))[0] + ".wav"
    )
    audio.export(wav_file_path, format="wav")
    print(f"[INFO] Conversion complete. Saved as '{wav_file_path}'")
    return wav_file_path


# 2. 오디오 파일을 청크로 나누기 (더 작은 단위로)
def split_audio(file_path, chunk_length_ms=8 * 60 * 1000):  # 8분으로 줄임
    print(f"[INFO] Splitting audio file into {chunk_length_ms // 60000} minute chunks...")
    audio = AudioSegment.from_file(file_path)
    chunks = []
    for i in range(0, len(audio), chunk_length_ms):
        chunks.append(audio[i : i + chunk_length_ms])
    print(f"[INFO] Created {len(chunks)} chunks")
    return chunks


# 3. Whisper로 전사 (타임스탬프 정보 활용)
def transcribe_audio_chunk(chunk, chunk_index):
    temp_file = f"temp_chunk_{chunk_index}.wav"
    print(f"[INFO] Transcribing chunk {chunk_index + 1}...")
    
    chunk.export(temp_file, format="wav")
    
    # Whisper에서 word-level timestamps 사용
    result = whisper_model.transcribe(
        temp_file, 
        word_timestamps=True,
        verbose=False
    )
    
    os.remove(temp_file)
    
    # Advanced formatting 적용
    formatted_text = advanced_format_transcription(result["text"])
    
    print(f"[INFO] Chunk {chunk_index + 1} completed")
    return formatted_text


# 4. 향상된 DOCX 저장
def save_text_to_docx(text, output_file):
    print(f"[INFO] Saving to '{output_file}'...")
    doc = Document()
    
    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = doc.styles['Normal'].font.size
    
    lines = text.split('\n')
    for line in lines:
        if line.strip():
            paragraph = doc.add_paragraph()
            
            # **볼드** 텍스트 처리
            if '**' in line:
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    else:
                        paragraph.add_run(part)
            else:
                paragraph.add_run(line)
        else:
            doc.add_paragraph()  # 빈 줄 유지
    
    doc.save(output_file)
    print(f"[INFO] Saved successfully!")


# 5. 메인 처리 함수
def process_large_audio(file_path, output_directory, create_bilingual=True):
    wav_file_path = convert_to_wav(file_path, output_directory)
    chunks = split_audio(wav_file_path)
    full_transcript = ""
    
    for index, chunk in enumerate(chunks):
        print(f"[INFO] Processing chunk {index + 1}/{len(chunks)}...")
        chunk_transcript = transcribe_audio_chunk(chunk, index)
        full_transcript += chunk_transcript + "\n\n"
    
    print(f"[INFO] Cleaning up temporary files...")
    os.remove(wav_file_path)
    
    # 전체 텍스트 후처리
    full_transcript = full_transcript.strip()
    
    # 이중 언어 형식 생성
    bilingual_text = None
    if create_bilingual:
        print("[INFO] Creating bilingual format...")
        bilingual_text = create_bilingual_format(full_transcript)
    
    return full_transcript, bilingual_text


# 메인 실행부
if __name__ == "__main__":
    directory = r"C:\Users\Jonathan\Documents\kbdocs\text_result"
    media_files = [f for f in os.listdir(directory) if f.endswith((".mp4", ".mp3"))]
    
    CREATE_BILINGUAL = True  # 이중언어 형식 생성 (영어 볼드 + 한국어)
    
    if not media_files:
        print("[WARNING] No media files found in the directory.")
    else:
        for media_file in media_files:
            audio_file_path = os.path.join(directory, media_file)
            print(f"\n[INFO] Processing: {media_file}")
            print("=" * 50)
            
            # 처리 실행
            transcript, bilingual = process_large_audio(
                audio_file_path, directory, CREATE_BILINGUAL
            )
            
            # 결과 저장
            base_name = os.path.splitext(media_file)[0]
            
            if bilingual:
                # 이중언어 파일 저장 (영어 볼드 + 한국어)
                bilingual_file = os.path.join(directory, f"{base_name}_bilingual.docx")
                save_text_to_docx(bilingual, bilingual_file)
                print(f"[INFO] Bilingual file saved: {bilingual_file}")
            else:
                # 영어만 저장
                transcript_file = os.path.join(directory, f"{base_name}_transcript.docx")
                save_text_to_docx(transcript, transcript_file)
                print(f"[INFO] Transcript saved: {transcript_file}")
            
            print(f"[INFO] Completed: {media_file}")
            print("=" * 50)